#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo file .docx "ĐIỂM TIN NGÀY d.M.yyyy" chứa các tin VỪA QUÉT ĐƯỢC trong lần publish này.
Cách xác định "tin mới của lần quét": diff DATA trong index.html (HEAD) với bản trước
(git show HEAD~1:index.html) — URL nào chưa có ở bản trước là tin của lần quét này.

⛔ BÁM CHẶT FORM MẪU CỦA CƠ QUAN — file `ĐTN_M_01.9.2026.docx` Huy gửi ngày 01/09/2026, chỉ
thị nguyên văn: *"từ bây giờ các kết quả quét tin xuất file docx phải theo form như file tao
đính kèm."* Đó là tên file MẪU, KHÔNG phải tên file script này xuất ra (`ten_file()` đặt).

04 mục, đánh số trong ngoặc đơn:
  (1) Đối ngoại Mỹ  -> usNews không khí tài, không Mali, có neo nước ngoài (`la_doi_ngoai_my`)
  (2) Nội bộ Mỹ     -> usNews không khí tài, không Mali, phần còn lại
  (3) Địa bàn Australia và Anh, Biển Đông -> worldNews neo được Úc/Biển Đông, chia 03 TIỂU MỤC
                       đậm không đánh số: Anh · Australia · Biển Đông (`tieu_muc_dia_ban`).
                       ⚠️ Tiểu mục "Anh" hiện LUÔN RỖNG — chủ đề 2 chưa mở phạm vi sang Anh,
                       cổng nạp `add_news.py::check_neo_chu_de_2` vẫn chặn tin thuần Anh.
  (4) KHCN-QS       -> usNews category "Công nghệ quân sự" + item tập trận/sự kiện mới
(Đã BỎ mục Mạng xã hội (X) và mục Mỹ – Mali — ngoài phạm vi / đi ở bản sáng.)

⛔ 05 điểm của form mẫu, mỗi điểm đều là chỗ dễ "sửa cho gọn" rồi lệch mẫu:
  - KHÔNG có dòng tiêu đề "ĐIỂM TIN NGÀY …" — vào thẳng mục (1).
  - Tin KHÔNG mở bằng gạch đầu dòng; mở bằng "Ngày d.M.yyyy," (ngày có số 0 dẫn, tháng không).
  - Link nằm CÙNG đoạn với nội dung, cách một dấu trắng — không xuống dòng riêng.
  - Mỗi đoạn: thụt dòng đầu 0,5"; giãn dòng CHÍNH XÁC 18pt (`line=360 exact`); cách đoạn 6pt.
  - Khổ A4, lề 1,0" cả bốn phía (bản cũ để 1,25" trái/phải).

⛔ MỤC 5 "Tin Jay Lâm gửi" ĐÃ BỎ HẲN 01/08/2026 — Huy đảo nguyên tắc: *"file của Jay Lâm gửi
chỉ là để so sánh xem có tin nào mày quét được mà bị trùng với tin trong file đó không thôi"*
· *"nếu có tin bị trùng với file Jay Lâm thì tự xoá khỏi tổng hợp tin đã quét đi và gửi file
word (trong đó không có tin nào từ Jay Lâm)"*. File Jay Lâm nay là **BỘ LỌC**: nó không đóng
góp dòng nào vào bản tin, chỉ dùng để bớt tin CỦA MÌNH mà anh ta đã đọc rồi.
Việc đối chiếu cần đọc hiểu theo SỰ KIỆN nên thuộc về PHIÊN QUÉT (có agent), không làm được
ở đây; phiên quét khai kết quả vào sổ `logs/trung-jaylam.json` qua
`scripts/tin_jaylam.py --ghi-loai`, và file này chỉ việc đọc sổ rồi bỏ tin — xem
`doc_url_trung_jaylam()` / `loc_bo_trung_jaylam()`. Vì thế file này KHÔNG còn chạm Supabase.

Định dạng khớp mẫu (đo bằng cách bóc `word/document.xml` của chính file mẫu):
  - Chữ: Times New Roman 14pt toàn bộ, khai ở CẢ style Normal lẫn từng run.
  - Đầu mục "(N) <tên>": căn đều, đậm, thụt dòng đầu 0,5", giãn dòng auto 276.
  - Tiểu mục (chỉ mục 3): "Anh"/"Australia"/"Biển Đông" — đậm, cùng khuôn đoạn với tin.
  - Mỗi tin: MỘT đoạn "Ngày d.M.yyyy, <nội dung>. <link>" (CHỈ summary — đã bỏ significance
    từ 27/07/2026), căn đều, thụt dòng đầu 0,5", giãn dòng exact 360, cách đoạn 6pt;
    link là hyperlink xanh gạch chân nằm cuối chính đoạn đó.
  - Một dòng trống sau mỗi mục.
  - Khổ A4; lề 1,0 inch cả bốn phía.
Xuất ra đường dẫn in ở stdout (dòng cuối "DOCX=<path>"). Rỗng (không có tin) -> in "DOCX=".
Tên file GỌI THEO BUỔI (chỉ thị Huy 28/07/2026): /tmp/Diem-tin-sang-som-5h-<ngày>.docx hoặc
/tmp/Diem-tin-toi-21h-<ngày>.docx — xem hàm `ten_file()`.

Chạy: python3 .github/scripts/make_docx.py
"""
import datetime, json, os, re, subprocess, sys, unicodedata, zoneinfo

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

FONT = "Times New Roman"
SIZE = 14  # pt — khớp mẫu

VN = zoneinfo.ZoneInfo("Asia/Ho_Chi_Minh")


def ten_file(gen, now=None):
    """Tên file .docx GỌI THEO BUỔI (chỉ thị Huy 28/07/2026): nhìn tên là biết bản nào.

      Diem-tin-sang-som-5h-<ngày>.docx  ·  Diem-tin-toi-21h-<ngày>.docx

    Ngưỡng buổi 14h giờ VN — CÙNG quy ước với `send_telegram.py:slot_label`,
    `send-email.js` và ô khoá `scripts/state.py`. Đổi lịch quét thì xem lại cả bốn nơi.

    ⚠️ ĐÂY LÀ NƠI DUY NHẤT ĐẶT TÊN FILE. Telegram (kênh gửi duy nhất hiện nay) hiện
    đúng basename của file trên đĩa, còn `send-email.js` lấy lại bằng `path.basename`
    thay vì tự ghép tên — hai bộ luật song song chắc chắn sẽ lệch, mà lệch âm thầm.
    """
    now = now or datetime.datetime.now(VN)
    buoi = "sang-som-5h" if now.hour < 14 else "toi-21h"
    return f"Diem-tin-{buoi}-{(gen or 'khong-ro-ngay').replace('/', '-')}.docx"


def extract_data(html):
    i = html.find("var DATA")
    if i < 0:
        raise ValueError('không thấy "var DATA"')
    start = html.find("{", i)
    depth = 0
    end = -1
    for k in range(start, len(html)):
        c = html[k]
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                end = k
                break
    if end < 0:
        raise ValueError("không đóng được object DATA")
    return json.loads(html[start:end + 1])


def prev_data():
    """DATA của index.html ở commit cha (HEAD~1). Lỗi -> None."""
    try:
        out = subprocess.run(
            ["git", "show", "HEAD~1:index.html"],
            capture_output=True, text=True, timeout=60,
        )
        if out.returncode != 0 or not out.stdout:
            return None
        return extract_data(out.stdout)
    except Exception:
        return None


def event_items(data):
    """Gom item con của exercises (Predator...) thành list phẳng.

    CHỈ lấy `exercises` — bản tin TỐI không đưa sự kiện ngoại giao (dipEvents do phiên
    SÁNG tạo, gửi qua notify-morning). Predator's Run gộp vào mục QS-KHCN.
    """
    items = []
    for grp in ("exercises",):
        for ev in data.get(grp, []) or []:
            ev_name = ev.get("name", "")
            for it in ev.get("items", []) or []:
                it = dict(it)
                it["_event"] = ev_name
                items.append(it)
    return items


def urls_of(items, key="sourceUrl"):
    return {it.get(key) for it in items if it.get(key)}


def diff_new(cur, prev, kind):
    """Trả list tin mới (có trong cur, không có trong prev). kind: usNews|worldNews|events."""
    if kind == "events":
        cur_list = event_items(cur)
    else:
        cur_list = cur.get(kind, []) or []  # worldNews / usNews

    if prev is None:
        # Không có bản trước -> fallback: lấy tin đưa lên hôm nay
        today = cur.get("generatedAt")
        return [it for it in cur_list
                if it.get("_addedDate") == today or it.get("date") == today]

    if kind == "events":
        prev_urls = urls_of(event_items(prev))
    else:
        prev_urls = urls_of(prev.get(kind, []) or [])

    return [it for it in cur_list if it.get("sourceUrl") and it.get("sourceUrl") not in prev_urls]


def today_items(cur, kind):
    """Toàn bộ tin đưa lên hôm nay (fallback khi diff rỗng)."""
    today = cur.get("generatedAt")
    lst = event_items(cur) if kind == "events" else (cur.get(kind, []) or [])
    return [it for it in lst if it.get("_addedDate") == today or it.get("date") == today]


def _khoa_tin(it):
    """Khoá nhận dạng một tin. `sourceUrl` là khoá chính; tin thiếu link (hay gặp ở mục
    tập trận) lùi về tiêu đề — dùng chuỗi rỗng làm khoá thì mọi tin thiếu link gộp thành
    một, tức bản tin nuốt mất tin mà không có tiếng kêu."""
    url = (it.get("sourceUrl") or "").strip()
    return url or ("T:" + str(it.get("title") or "") + "|" + str(it.get("summary") or "")[:60])


def pick_items(cur, prev, kind):
    """Tin của bản tin hôm nay = HỢP của (mới so với commit cha) và (_addedDate/date == generatedAt).

    ⚠️ Sự cố 25/07/2026 — docx chỉ có 3/15 tin: phiên quét commit `index.html` HAI lần
    (một commit `log: checkpoint ...` giữa chừng, rồi commit `Cap nhat ban tin` cuối).
    Chỉ commit cuối gửi email, nhưng lúc đó HEAD~1 đã chứa sẵn 12 tin nạp ở lô đầu nên
    `diff_new` chỉ còn 3 tin — và fallback cũ chỉ chạy khi diff RỖNG HẲN nên không cứu.
    Ngược lại, chỉ dùng `today_items` lại hụt tin của lô neo ngày cũ (`_addedDate` lệch
    `generatedAt`, xem "HAI BẪY khi lô tin trải QUÁ 2 NGÀY" trong CLAUDE.md).
    → Lấy HỢP để chắc cả hai chiều. Giữ nguyên thứ tự trong mảng gốc, không trùng lặp.
    """
    lst = event_items(cur) if kind == "events" else (cur.get(kind, []) or [])
    today = cur.get("generatedAt")
    new_urls = urls_of(diff_new(cur, prev, kind))
    out = []
    for it in lst:
        is_today = it.get("_addedDate") == today or it.get("date") == today
        url = it.get("sourceUrl")
        if is_today or (url and url in new_urls):
            out.append(it)
    return out


# Từ khoá nhận tin chủ đề 4 (Mỹ–Mali/Sahel) — cũng nằm trong `usNews` với category
# "Chính trị", nên phải tách ra khỏi mục "Nội bộ Mỹ". Viết KHÔNG DẤU vì so sau khi bỏ dấu.
MALI_KEYS = ("mali", "jnim", "bamako", "sahel", "azawad", "niger", "burkina",
             "africa corps", "chau phi", "sahen")
# Region được coi là "trong nước Mỹ". Rỗng cũng tính — xem chú thích trong is_noibo_my.
REGION_NOI_BO = ("", "Bắc Mỹ", "Châu Mỹ")


def _khong_dau(s):
    """Bỏ dấu tiếng Việt để so từ khoá. Cố ý viết tại chỗ thay vì import từ
    `scripts/tra_cuu_tin.py`: file này nằm trên đường đi của email hằng ngày, một lỗi
    import chéo thư mục là mất file .docx của cả bản tin."""
    s = unicodedata.normalize("NFD", str(s or "").lower())
    return "".join(c for c in s if unicodedata.category(c) != "Mn").replace("đ", "d")


# So theo BIÊN TỪ, không so chuỗi con — vá 26/08/2026. Đo trên kho thật hôm đó: tin *"Hải quân
# Mỹ công bố tên lửa không đối không tầm xa AIM-424 Malice"* bị gán chủ đề Mỹ–Mali vì chuỗi
# "mali" nằm trong chữ "Malice", rồi bị loại khỏi .docx bản tối theo chỉ thị 05/08/2026 — mất
# một tin công nghệ quân sự mà không dấu hiệu nào, .docx vẫn đủ mục. Cùng lối đó, "niger" khớp
# "Nigeria". `scripts/topics.py` vốn đã so bằng biên từ; chỗ này là bản chép lệch còn sót.
_RE_MALI = tuple(re.compile(r"(?<!\w)" + re.escape(k) + r"(?!\w)") for k in MALI_KEYS)


def la_tin_mali(it):
    kho = _khong_dau(" ".join(str(it.get(k, "")) for k in
                              ("title", "summary", "region", "significance")))
    return any(p.search(kho) for p in _RE_MALI)


def _kho_chu(it):
    return " ".join(str(it.get(k, "")) for k in ("title", "summary", "region", "significance"))


# Bảng neo của chủ đề 2 nằm ở `scripts/topics.py` — MỘT hàm kiểm tra duy nhất, dùng chung
# với cổng nạp `add_news.py`. Chép bảng sang đây thì hai bản tách nhánh ở lần vá sau mà
# không ai thấy.
#
# ⚠️ ĐÁNH ĐỔI CÓ CHỦ Ý, khác hẳn `_khong_dau` ở trên (hàm đó chép tại chỗ vì nó chỉ là
# 2 dòng, chép không sinh ra nguồn sự thật thứ hai). Ở đây phần chép sẽ là một BẢNG TỪ
# KHOÁ ~50 mục — thứ chắc chắn được sửa tiếp và chắc chắn lệch. Nên chấp nhận import chéo
# thư mục, và CỐ Ý ĐỂ NÉM LỖI thay vì `try/except` cho êm: import hỏng thì .docx không
# sinh ra và CI đỏ ngay, còn `except: lambda _: True` thì mục 2 lặng lẽ trở lại làm cái
# thùng chứa — đúng thứ đang đi vá. File mất thì có tiếng kêu; file sai nội dung thì không.
sys.path.insert(0, os.path.join(
    os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "scripts"))
from topics import neo_uc_bien_dong, neo_uc, neo_anh  # noqa: E402


def la_uc_bien_dong(it):
    """Tin có tự neo được vào Úc hoặc Biển Đông không? (điều kiện vào mục 2)

    ⚠️ VÁ 01/08/2026 — Huy bắt: *"hàn quốc liên quan đ gì đến biển đông và Úc mà cứ cho
    vào???"*. Trước đây mục 2 được định nghĩa là **MỌI `worldNews` trừ Mali**, tức một cái
    thùng: tin thế giới nào lọt qua tầng quét cũng tự động được dán nhãn "Úc và Biển Đông",
    nên tên mục là một lời hứa còn nội dung là phần dư. Bản tối 01/08 có 04 tin thì 03 sai
    (Nhật phóng Tomahawk từ JS Chokai · Trung Quốc phóng YJ-20 · Hàn ký 7,8 nghìn tỷ won
    với Hanwha Ocean) — cả ba không dính Úc, không dính Biển Đông.

    Đây là CÙNG con lỗi đã vá hai lần ở `is_noibo_my`, chỉ khác mảng: mục nào được định
    nghĩa bằng "phần còn lại" thì mọi phân loại thiếu sót đều đổ vào nó. Nay cả bốn mục
    đều định nghĩa DƯƠNG.

    Tin rớt KHÔNG biến mất — lưới an toàn ở `build_sections` gom về mục 1 kèm cảnh báo, vì
    mất tin tệ hơn xếp nhầm mục. Chính cảnh báo đó là tín hiệu cho biết tầng quét đã lọt.
    """
    return neo_uc_bien_dong(_kho_chu(it))


# Category của tin thuộc mục QS-KHCN. Đây là danh sách DƯƠNG — xem chú thích is_noibo_my.
CATEGORY_QSKHCN = ("Công nghệ quân sự",)


def la_qs_khcn(it):
    return (it.get("category") or "") in CATEGORY_QSKHCN


def is_noibo_my(it):
    """Nội bộ Mỹ = tin `usNews` KHÔNG phải khí tài, KHÔNG phải chuyện Mali/Sahel.

    ⚠️ VÁ LẦN 1 (27/07/2026) — trước đây điều kiện là `region == "Bắc Mỹ"`, và mục 1 của
    file .docx vì thế LUÔN RỖNG: mọi tin `usNews` nạp gần đây đều không có `region` (đếm
    thật: 9/9 tin đều `region: None`). Vì vậy region RỖNG vẫn được tính là nội bộ Mỹ — đó
    là trạng thái bình thường của dữ liệu. Đừng "siết cho chặt" bằng cách bắt buộc có
    region: làm vậy là tái lập đúng con bug đó.

    ⚠️ VÁ LẦN 2 (27/07/2026, sau khi Huy bắt lỗi tin Mali lòi ra mục QS-KHCN) — điều kiện
    cũ `category == "Chính trị"` cũng sai, chỉ theo hướng ngược lại: nó BỎ SÓT tin **Kinh
    tế** và **Ngoại giao**, mà theo 5 nhóm Nội bộ Mỹ Huy chốt thì nhóm 4 chính là *"hoạt
    động kinh tế Mỹ + hoạt động khác của các bộ và Nhà Trắng"*. Đếm thật trên DATA: 30 tin
    `Kinh tế` + 21 tin `Ngoại giao` đang bị dồn xuống mục khí tài — thực tế lọt vào bản
    27/07: "Mỹ áp gói thuế quan mới 10-12,5% lên khoảng 60 đối tác thương mại".

    GỐC của cả hai lần vá là cùng một thứ: **QS-KHCN từng được định nghĩa là "mọi usNews
    còn lại" — một cái thùng rác**, nên mọi phân loại thiếu sót đều đổ vào đó. Nay CẢ HAI
    mục đều định nghĩa DƯƠNG: QS-KHCN = category "Công nghệ quân sự"; Nội bộ Mỹ = phần còn
    lại KHÔNG phải khí tài và KHÔNG phải Mali. Thêm category mới thì cân nhắc nó thuộc mục
    nào, đừng để rơi tự do.
    """
    if la_qs_khcn(it):
        return False
    if la_tin_mali(it):
        return False
    return (it.get("region") or "") in REGION_NOI_BO


# ═══════════════════ TÊN MỤC — khai MỘT chỗ, mọi nơi khác tham chiếu ═══════════════════
# ⛔ Bộ test PHẢI import các hằng này, CẤM chép chuỗi tên mục vào test. Chép chuỗi thì lần
# đổi tên mục sau đây bộ test đỏ vì tên chứ không vì hành vi, và người sửa sẽ sửa test cho
# hết đỏ thay vì soi hành vi — đúng lối làm cổng kiểm mất răng.
MUC_DOI_NGOAI = "Đối ngoại Mỹ"
MUC_NOI_BO = "Nội bộ Mỹ"
MUC_DIA_BAN = "Địa bàn Australia và Anh, Biển Đông"
MUC_KHCN = "KHCN-QS"

# Tiểu mục của mục địa bàn — mẫu Huy gửi 01/09/2026 in chúng thành 03 dòng đậm không đánh số.
TM_ANH = "Anh"
TM_UC = "Australia"
TM_BIEN_DONG = "Biển Đông"
THU_TU_TIEU_MUC = (TM_ANH, TM_UC, TM_BIEN_DONG)

# Khoá gắn tạm lên từng tin của mục địa bàn để `main()` biết in nó dưới tiểu mục nào.
# Cố ý gắn lên chính dict tin (bản sao) thay vì đổi kiểu trả về của `build_sections`:
# 03 bộ test hiện có đọc `[(tên_mục, [tin])]`, đổi kiểu là chúng vỡ vì hình dạng dữ liệu
# chứ không vì hành vi phân mục.
KHOA_TIEU_MUC = "_tieuMuc"

# ⛔ MỌI tin trong file .docx nay đều mở bằng "Ngày d.M.yyyy," (form mẫu 01/09/2026), nên
# KHÔNG còn mục nào được ghi ngày riêng. Hằng cũ `MUC_GHI_NGAY` đã bỏ — trước đây chỉ mục
# QS-KHCN ghi ngày vì chỉ nó được nới khung 3 ngày; nay khung ngày không đổi, chỉ cách in đổi.


# Từ khoá neo tin `usNews` vào mục "Đối ngoại Mỹ" (tách khỏi "Nội bộ Mỹ" theo mẫu
# 01/09/2026). Viết KHÔNG DẤU vì so sau khi bỏ dấu, và so theo BIÊN TỪ.
#
# ⚠️ Bảng này định nghĩa DƯƠNG cho mục 1, còn mục 2 "Nội bộ Mỹ" là phần `usNews` còn lại.
# Đó là ngoại lệ có chủ ý so với nguyên tắc "mọi mục định nghĩa dương" ở `is_noibo_my`:
# hai mục này chia đôi CÙNG một tập (`usNews` không phải khí tài, không phải Mali), nên
# phần còn lại ở đây là một tập ĐÃ ĐÓNG, không phải cái thùng hứng mọi thứ rơi.
#
# ⛔ KHÔNG thêm tên hãng tin ("reuters", "bbc") — tin nội bộ nào cũng dẫn nguồn nước ngoài.
# ⛔ KHÔNG thêm "trung dong"/"chau au" trần: quá rộng, và tin nội bộ hay nhắc để lấy bối cảnh.
#
# ⛔ TÊN NƯỚC TIẾNG VIỆT MỘT ÂM TIẾT PHẢI ĐI KÈM "nuoc"/"tong thong"/… — CẤM để trần.
#    Đo thật 01/09/2026 ngay lượt chạy thử đầu: neo `"duc"` (Đức) khớp chữ **"tình dục"** của
#    tin *"Hạ viện Mỹ thông qua Đạo luật Kayleigh"*, đẩy một tin nội bộ thuần sang mục
#    "Đối ngoại Mỹ". Cùng lối đó `"phap"` khớp "biện pháp"/"tư pháp"/"pháp luật" — thứ có
#    trong gần như MỌI tin lập pháp Mỹ, và `"nga"` khớp "ngã". Biên từ KHÔNG cứu được ở đây
#    vì sau khi bỏ dấu chúng là từ trọn vẹn. Đây đúng lớp lỗi "mali khớp Malice" đã vá
#    26/08/2026, chỉ khác bảng.
NEO_DOI_NGOAI = (
    # -- nước & vùng lãnh thổ hay dính chính sách đối ngoại Mỹ
    "iran", "nuoc nga", "lien bang nga", "tong thong nga", "russia", "russian",
    "moscow", "putin", "ukraine", "kyiv",
    "trung quoc", "china", "chinese", "bac kinh", "beijing",
    "israel", "gaza", "bo tay", "west bank", "palestine",
    "trieu tien", "north korea", "binh nhuong", "pyongyang",
    "venezuela", "cuba", "mexico", "canada", "panama", "greenland",
    "nuoc anh", "vuong quoc anh", "united kingdom", "britain", "british", "london",
    "argentina", "falkland",
    "nuoc phap", "france", "french", "paris",
    "nuoc duc", "germany", "german", "berlin", "ba lan", "poland",
    "nhat ban", "japan", "han quoc", "south korea", "dai loan", "taiwan",
    "an do", "india", "pakistan", "afghanistan", "iraq", "syria", "yemen", "houthi",
    "saudi", "qatar", "uae", "tho nhi ky", "turkey", "ai cap", "egypt",
    "brazil", "colombia",
    # -- cơ chế đa phương
    "nato", "g20", "g7", "lien hop quoc", "united nations", "lien minh chau au",
    "european union", "opec", "asean", "imf", "world bank",
    # -- công cụ đối ngoại
    "trung phat", "sanction", "cam van", "embargo", "thue quan", "tariff",
    "vien tro nuoc ngoai", "foreign aid", "hiep dinh", "dam phan thuong mai",
    "dai su", "ambassador", "ngoai truong", "secretary of state",
)
_RE_DOI_NGOAI = tuple(re.compile(r"(?<!\w)" + re.escape(k) + r"(?!\w)")
                      for k in NEO_DOI_NGOAI)


def la_doi_ngoai_my(it):
    """Tin `usNews` này nói chuyện Mỹ với BÊN NGOÀI hay chuyện trong nước?

    Hai đường vào: category `Ngoại giao` (tầng quét đã phân), hoặc câu chữ tự neo được vào
    một nước/cơ chế/công cụ đối ngoại. Đường thứ hai là cần thiết vì phần lớn tin đối ngoại
    trong kho mang category `Chính trị` hoặc `Kinh tế` — đo trên kho 01/09/2026: tin trừng
    phạt Iran của Bộ Tài chính Mỹ mang cat `Kinh tế`, tin G20 mời Nga cũng vậy.
    """
    if (it.get("category") or "") == "Ngoại giao":
        return True
    kho = _khong_dau(_kho_chu(it))
    return any(p.search(kho) for p in _RE_DOI_NGOAI)


def tieu_muc_dia_ban(it):
    """Tin của mục địa bàn thuộc tiểu mục nào: Anh · Australia · Biển Đông.

    ⚠️ THỨ TỰ GIÀNH KHÁC THỨ TỰ IN. In theo mẫu: Anh · Australia · Biển Đông
    (`THU_TU_TIEU_MUC`). Giành thì **Úc trước Anh**, vì tin dính CẢ HAI nước gần như luôn là
    AUKUS hoặc tập trận do Úc chủ trì có Anh dự — thuộc về Úc, chủ đề gốc của mảng; Anh mới
    là phần thêm 01/09/2026. Đảo lại thì mọi tin AUKUS rơi xuống tiểu mục "Anh".
    Tin Anh thuần vẫn về đúng chỗ: *"HMS Tamar thăm Căn cứ Ream/Campuchia"* không có neo Úc
    nào, nên nó qua nhánh Úc rồi rơi đúng vào nhánh Anh.

    Không khớp Úc lẫn Anh thì về "Biển Đông" — đây là phần còn lại của một tập ĐÃ ĐÓNG
    (tin đã qua cổng `la_uc_bien_dong`), không phải thùng rác.
    """
    kho = _kho_chu(it)
    if neo_uc(kho):
        return TM_UC
    if neo_anh(kho):
        return TM_ANH
    return TM_BIEN_DONG


def build_sections(us, world, events):
    """Chia thành 3 mục của bản tin (trước 05/08/2026 là 4 — mục Mali đã BỎ).

    ⛔ **MỤC "Mỹ – Mali" BỎ KHỎI FILE WORD 05/08/2026 (chỉ thị Huy)**, nguyên văn: *"bỏ mục
    Mali trong file word gửi tele hàng ngày. Thêm mục Mali vào kết quả phần quét tập trận và
    thinktank."* Tin Mali **KHÔNG mất**: nó vẫn được quét, vẫn nạp vào `DATA`, vẫn lên web, và
    nay đi ở **bản sáng 🎖️ Sự kiện & Tập trận** (`send-morning-email.js::diffMali`) cùng chỗ
    với tập trận và think-tank.

    ⚠️ **PHÉP LỌC MALI PHẢI GIỮ NGUYÊN, CHỈ BỎ MỤC.** Đây là chỗ dễ vá sai nhất: xoá luôn
    `mali`/`mali_urls` thì tin Sahel hết bị tách ra, và mục 1 "Nội bộ Mỹ" lại hứng chúng —
    đúng con lỗi Huy bắt 27/07/2026 (*"đang tin khcn-qs tự nhiên thấy lòi ra tin Mali"*), chỉ
    khác chỗ rơi. Vì vậy `mali_urls` VẪN nằm trong `da_xep` để lưới an toàn không dồn chúng
    về mục 1, và hàm in một dòng KÊU mỗi lần bỏ — im lặng ở đây thì không ai biết bản tin đã
    rụng mấy tin.

    ⚠️ ĐÃ VÁ 27/07/2026 — Huy bắt lỗi: *"đang tin khcn-qs tự nhiên thấy lòi ra tin Mali, và
    chẳng thấy mục mali đâu"*. Trước đây hàm này chỉ dựng 3 mục và mục QS-KHCN được định
    nghĩa là "MỌI usNews còn lại", nên tin Mỹ–Mali (một trong 5 chủ đề, có mục riêng trên
    web) bị dồn vào đó nằm lẫn giữa tin khí tài — người đọc vừa thấy lạc lõng vừa mất hẳn
    một chủ đề. Thực tế lọt vào bản 27/07: "Al Jazeera phân tích liên minh JNIM", "Niger
    Abdourahamane Tiani… Mali, Burkina".

    Hai điểm phải giữ:
    - Lọc Mali từ CẢ `us` LẪN `world`: tin Sahel nằm ở mảng nào cũng có thể, và trước đây
      `world` được đổ nguyên vào "Úc và Biển Đông" nên tin Mali trong `world` sẽ lọt vào mục
      Biển Đông — đúng cùng một con lỗi, chỉ khác chỗ.
    - Ba nhánh phải LOẠI TRỪ NHAU, nếu không một tin sẽ in hai lần ở hai mục.
    - Và phải PHỦ HẾT: từ khi QS-KHCN thôi làm "thùng rác hứng phần còn lại", một tin không
      khớp nhánh nào sẽ BIẾN MẤT khỏi file mà không báo gì. Lưới cuối bên dưới gom phần rơi
      về mục 1 và in cảnh báo — mất tin tệ hơn nhiều so với xếp nhầm mục.

    Diễn biến tập trận vẫn nằm trong QS-KHCN qua `events` (bản tin mẫu để vậy, Huy không đụng).
    """
    mali = [it for it in us + world if la_tin_mali(it)]
    mali_urls = urls_of(mali)
    if mali:
        # KÊU, không im. Tin bị bỏ khỏi .docx phải soi ngược được — cùng nguyên tắc với
        # `loc_bo_trung_jaylam`. Đây KHÔNG phải lỗi: đúng chỉ thị Huy 05/08/2026.
        print(f"ℹ️  {len(mali)} tin Mỹ–Mali KHÔNG vào .docx (chỉ thị Huy 05/08/2026) — chúng "
              f"đi ở bản sáng 🎖️ Sự kiện & Tập trận: "
              + " | ".join((it.get("title") or "")[:45] for it in mali[:5]), file=sys.stderr)

    def khong_phai_mali(it):
        return it.get("sourceUrl") not in mali_urls

    my = [it for it in us if is_noibo_my(it)]           # usNews không khí tài, không Mali
    sec1 = [it for it in my if la_doi_ngoai_my(it)]                   # 1. Đối ngoại Mỹ
    sec2 = [it for it in my if not la_doi_ngoai_my(it)]               # 2. Nội bộ Mỹ
    sec3 = [dict(it, **{KHOA_TIEU_MUC: tieu_muc_dia_ban(it)})         # 3. Địa bàn (03 tiểu mục)
            for it in world if khong_phai_mali(it) and la_uc_bien_dong(it)]
    sec4 = [it for it in us                                           # 4. KHCN-QS (+ tập trận)
            if la_qs_khcn(it) and khong_phai_mali(it)]

    # LƯỚI AN TOÀN — không được để tin nào rơi ra ngoài mọi mục.
    da_xep = urls_of(sec1) | urls_of(sec2) | urls_of(sec3) | urls_of(sec4) | mali_urls
    roi = [it for it in us + world
           if it.get("sourceUrl") and it.get("sourceUrl") not in da_xep]
    if roi:
        # Tách riêng nhóm tin THẾ GIỚI rơi vì không neo được vào Úc/Biển Đông (siết
        # 01/08/2026). Nhóm này khác hẳn nhóm rơi vì thiếu category: nó nói rằng TẦNG QUÉT
        # đã nạp tin ngoài phạm vi 5 chủ đề, tức phải sửa ở `add_news.py`/phiên quét chứ
        # không phải sửa phân loại ở đây. Gộp chung một dòng cảnh báo thì hai nguyên nhân
        # khác nhau ra cùng một câu chữ, và người đọc sẽ đi sửa nhầm chỗ.
        world_urls = urls_of(world)
        lac_dia_ban = [it for it in roi if it.get("sourceUrl") in world_urls]
        if lac_dia_ban:
            print(f"⚠️  {len(lac_dia_ban)} tin worldNews KHÔNG neo được vào Úc/Biển Đông -> "
                  f"tạm dồn vào {MUC_NOI_BO!r} để không mất tin. Đây là lỗi TẦNG QUÉT, "
                  f"không phải lỗi phân loại: "
                  + " | ".join((it.get("title") or "")[:45] for it in lac_dia_ban[:5]),
                  file=sys.stderr)
        con_lai = [it for it in roi if it.get("sourceUrl") not in world_urls]
        if con_lai:
            print(f"⚠️  {len(con_lai)} tin không khớp mục nào -> dồn vào {MUC_NOI_BO!r}. "
                  f"Xem lại phân loại: "
                  + " | ".join(f"[{it.get('category')}] {(it.get('title') or '')[:45]}"
                               for it in con_lai[:5]), file=sys.stderr)
        # ⛔ LƯỚI DỒN VÀO "Nội bộ Mỹ", TUYỆT ĐỐI KHÔNG DỒN VÀO MỤC ĐỊA BÀN — thử đổi lúc dựng
        # form mới 01/09/2026 và ca 11 của `tests/test-cong-uc-bien-dong.py` bắt ngay. Lý do
        # ca đó tồn tại (Huy bắt 01/08/2026): mục địa bàn từng là cái THÙNG, tin thế giới nào
        # lọt tầng quét cũng được dán nhãn "Úc và Biển Đông". Dồn lưới vào đúng mục ấy là
        # dựng lại cái thùng, chỉ khác là có thêm dòng kêu — mà dòng kêu nằm ở stderr của CI,
        # còn nhãn sai thì nằm trong bản tin Huy đọc. Tin thế giới nằm dưới nhãn "Nội bộ Mỹ"
        # trông CHƯỚNG, và chính chỗ chướng đó là tín hiệu để đi sửa tầng quét.
        sec2 = sec2 + roi

    # ⛔ KHÔNG thêm lại `("Mỹ – Mali", mali)` vào đây — xem docstring. Tin Mali đi ở bản sáng.
    return [
        (MUC_DOI_NGOAI, sec1),
        (MUC_NOI_BO, sec2),
        (MUC_DIA_BAN, sec3),
        (MUC_KHCN, sec4 + list(events)),
    ]


# ---------- docx helpers ----------
def set_font(run, size=SIZE, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), FONT)
    rpr.append(rfonts)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(SIZE * 2)); rpr.append(sz)  # 14pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0000FF"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def item_body(it):
    """Nội dung tin trong .docx: CHỈ `summary` (fallback `title` nếu thiếu summary).

    ⚠️ ĐÃ BỎ `significance` khỏi đây (chỉ thị Huy 27/07/2026: "ở nội dung tóm tắt tin trong
    file docx thì bỏ những câu đánh giá ở cuối"). Trước đây ghép `summary + significance`
    thành một đoạn, nên mỗi tin đều kết bằng một câu bình luận kiểu "cho thấy...", "phản
    ánh..." — đọc rất sáo và không phải thứ Huy cần trong bản tin.
    `significance` VẪN được giữ trong `DATA` và vẫn hiển thị trên web — chỉ không đưa vào
    file Word gửi kèm email. Đừng "dọn cho gọn" bằng cách xoá field này khỏi guardrail.
    """
    if it.get("summary"):
        return it["summary"].strip()
    if it.get("title"):
        return it["title"].strip()
    return ""


def ngay_ngan(s):
    """'2026-07-24' -> '24/07'. Không parse được thì trả nguyên chuỗi (còn hơn nuốt mất).

    Giữ lại dù form 01/09/2026 không còn dùng: `send_telegram.py` và bộ test gọi nó.
    """
    s = str(s or "").strip()
    m = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", s)
    return f"{m.group(3)}/{m.group(2)}" if m else s


def ngay_form(s):
    """'2026-09-01' -> '01.9.2026' — ĐÚNG khuôn mẫu Huy gửi 01/09/2026.

    ⚠️ Ngày GIỮ số 0 dẫn, tháng KHÔNG — đo trên mẫu: `Ngày 01.9.2026` và `Ngày 31.8.2026`.
    Đừng "cho đều" thành `01.09.2026`: đó là khuôn khác, và mẫu là thứ đang phải bám.
    Không parse được thì trả rỗng — thà không có tiền tố ngày còn hơn in ra một chuỗi lạ.
    """
    m = re.fullmatch(r"(\d{4})-(\d{2})-(\d{2})", str(s or "").strip())
    return f"{m.group(3)}.{int(m.group(2))}.{m.group(1)}" if m else ""


# Từ mở đầu câu phải HẠ chữ hoa khi bị đẩy vào sau "Ngày d.M.yyyy, ". Chỉ gồm trạng ngữ và
# từ chức năng — CỐ Ý không đụng tới chức danh ("Bộ trưởng", "Tổng thống", "Thủ tướng"):
# mẫu 01/09/2026 giữ hoa chúng ngay sau dấu phẩy (*"Ngày 31.8.2026, Bộ trưởng Tài chính…"*),
# đúng lối văn bản hành chính Việt Nam.
# ⛔ CẤM đưa vào đây những chữ vừa mở câu vừa mở CHỨC DANH hoặc TÊN CƠ QUAN — đo trên chính
#    mẫu: *"Đại tướng Dan Caine"* · *"Quân đội Mỹ và Indonesia"* · *"Lực lượng Không gian Mỹ"*
#    · *"Chính phủ Anh"* đều viết HOA ngay sau dấu phẩy. Bốn chữ "đại", "quân", "lực",
#    "chính" vì thế đã bị gỡ khỏi bảng sau lượt chạy thử 01/09/2026.
TU_HA_CHU_DAU = frozenset((
    "tại", "theo", "trong", "trước", "sau", "trên", "dưới", "giữa", "ngay",
    "một", "hai", "ba", "nhiều", "các", "những", "hàng", "khoảng", "gần", "hơn",
    "truyền", "khảo", "cuộc", "hoạt", "phát", "bang", "tàu", "máy", "đoàn",
    "phía", "phe", "số", "tỷ", "tình", "thị", "giá", "dự", "nhóm", "hãng",
    "báo", "tin", "nguồn", "cả", "toàn", "việc", "phiên", "kỳ", "vòng", "đợt",
))


def _noi_sau_ngay(body):
    """Hạ chữ đầu của `body` khi nó sắp đứng sau 'Ngày d.M.yyyy, '.

    Chỉ hạ khi từ đầu nằm trong `TU_HA_CHU_DAU`. Danh sách DƯƠNG chứ không phải phép đoán
    "chữ nào không phải tên riêng": đoán sai theo chiều hạ nhầm là viết thường tên người,
    tên nước — lỗi đập vào mắt ngay dòng đầu bản tin, mà không cổng nào bắt được.
    """
    tu = body.split(" ", 1)[0]
    return body[0].lower() + body[1:] if tu.lower() in TU_HA_CHU_DAU else body


# Tiền tố ngày mà chính tầng quét đã tự viết vào `summary` ("Ngày 31/8/2026, …"). Phải cắt
# rồi ghép lại theo khuôn mẫu, nếu không bản tin ra "Ngày 31.8.2026, Ngày 31/8/2026, …".
_RE_TIEN_TO_NGAY = re.compile(
    r"^\s*ngày\s+\d{1,2}\s*[./-]\s*\d{1,2}\s*[./-]\s*\d{4}\s*[,:]?\s*", re.IGNORECASE)


def than_tin(it):
    """Thân đoạn tin theo form mẫu: 'Ngày d.M.yyyy, <nội dung>'.

    Ngày lấy từ trường `date` của tin — KHÔNG lấy ngày mà tầng quét viết trong `summary`:
    `date` là con số đã qua cổng `scripts/ngay_that.py` (đối chiếu metadata trang gốc), còn
    ngày trong câu chữ là thứ agent tự gõ. Hai nguồn lệch nhau thì tin `date`.
    Không có `date` hợp lệ thì in thân tin trần, không bịa ngày.
    """
    body = _RE_TIEN_TO_NGAY.sub("", item_body(it)).strip()
    d = ngay_form(it.get("date"))
    return f"Ngày {d}, {_noi_sau_ngay(body)}" if (d and body) else body


def _dinh_dang_doan(p, dam=False):
    """Khuôn đoạn của mẫu: thụt dòng đầu 0,5\"; giãn dòng CHÍNH XÁC 18pt; cách đoạn 6pt.

    `w:line=360 w:lineRule="exact"` — mẫu dùng `exact`, không phải `auto`. Đổi sang `auto`
    là mỗi trang co lại vài dòng và bản in lệch hẳn so với bản của cơ quan.
    """
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.first_line_indent = Inches(0.5)
    pf.space_after = Pt(6) if not dam else Pt(0)
    # ⚠️ KHÔNG đặt `space_before` — mẫu không khai `w:before`, và python-docx ghi ra
    # `w:before="0"` khi được gán, tức thêm một thuộc tính mẫu không có.
    ppr = p._p.get_or_add_pPr()
    spc = ppr.find(qn("w:spacing"))
    if spc is None:
        spc = OxmlElement("w:spacing")
        ppr.append(spc)
    spc.set(qn("w:line"), "360")
    spc.set(qn("w:lineRule"), "exact")
    return p


def add_dau_muc(doc, chu):
    """Đầu mục '(N) Tên mục' — đậm, thụt dòng đầu, giãn dòng 1,15 (auto 276) như mẫu."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    ppr = p._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:line"), "276")
    spc.set(qn("w:lineRule"), "auto")
    ppr.append(spc)
    set_font(p.add_run(chu), size=SIZE, bold=True)
    return p


def add_tieu_muc(doc, chu):
    """Tiểu mục 'Anh' / 'Australia' / 'Biển Đông' — đậm, cùng khuôn đoạn với tin."""
    p = _dinh_dang_doan(doc.add_paragraph())
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(chu), size=SIZE, bold=True)
    return p


def add_dong_trong(doc):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    ppr = p._p.get_or_add_pPr()
    spc = OxmlElement("w:spacing")
    spc.set(qn("w:line"), "276")
    spc.set(qn("w:lineRule"), "auto")
    ppr.append(spc)
    return p


def add_item(doc, it):
    """Một tin = MỘT đoạn duy nhất: 'Ngày d.M.yyyy, <nội dung>. <link>'.

    ⚠️ Link nằm CÙNG đoạn với nội dung, cách một dấu trắng — không xuống dòng riêng như bản
    trước 01/09/2026. Đây là điểm dễ "sửa cho gọn" nhất và cũng là điểm mẫu khác rõ nhất:
    tách link ra dòng riêng thì mỗi tin chiếm gấp đôi số dòng và bản 4 trang phình thành 7.
    """
    body = than_tin(it)
    p = _dinh_dang_doan(doc.add_paragraph())
    if body:
        if not body.rstrip().endswith((".", "!", "?", ":")):
            body = body.rstrip() + "."
        set_font(p.add_run(body + " "), size=SIZE)
    url = it.get("sourceUrl")
    if url:
        add_hyperlink(p, url, url)


def loc_chua_gui(items):
    """Bỏ tin ĐÃ nằm trong một bản tin đã gửi trước đó (sổ logs/da-gui-email.json).

    Chỉ thị Huy 27/07/2026: bản tin TỐI phải "loại cả những tin đã quét lúc 4h 5h sáng" —
    tức chỉ gồm tin thu được SAU lần gửi trước. Không có sổ (chạy lần đầu, file lỗi) thì
    trả nguyên danh sách: thà gửi trùng còn hơn gửi rỗng.
    """
    try:
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from so_da_gui import url_da_gui
        da_gui = url_da_gui()
    except Exception as e:                  # noqa: BLE001
        print(f"Không đọc được sổ đã gửi ({e}) — giữ nguyên toàn bộ tin.", file=sys.stderr)
        return items
    if not da_gui:
        return items
    out = [it for it in items if it.get("sourceUrl") not in da_gui]
    if len(out) != len(items):
        print(f"Sổ đã gửi: bỏ {len(items) - len(out)} tin đã có trong bản tin trước.",
              file=sys.stderr)
    return out


def _url_ca_sang(now):
    """URL đã gửi ở ca SÁNG cùng ngày, hoặc tập RỖNG nếu không áp được.

    Một đường đọc sổ duy nhất cho mọi nơi cần biết "tin nào đã gửi ca sáng" — hai nơi tự đọc
    sổ riêng thì chắc chắn lệch nhau, mà lệch âm thầm (mục 14 CLAUDE.md toàn cục). Người gọi
    thứ hai (`loc_jaylam_ca_sang`) đã bỏ cùng mục 5 ngày 01/08/2026; nay chỉ còn
    `loc_bo_tin_ca_sang`, nhưng vẫn giữ hàm tách riêng để lớp sau khỏi tự dựng lại.

    Rỗng nghĩa là KHÔNG lọc gì: bản sáng (tin vừa nạp, chưa từng gửi) · sổ chưa có dòng nào ·
    đọc sổ hỏng. Hướng lệch có chủ ý là LẶP một bản tin, không phải MẤT tin.
    """
    if not la_buoi_toi(now):
        return set()
    try:
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from so_da_gui import url_da_gui_buoi
        return url_da_gui_buoi("sang", now.strftime("%Y-%m-%d")) or set()
    except Exception as e:                  # noqa: BLE001
        print(f"Không đọc được sổ ca sáng ({e}) — giữ nguyên toàn bộ tin.", file=sys.stderr)
        return set()


def _doc_url_buoi(buoi, ngay):
    """Tập URL đã gửi ở ĐÚNG một buổi, ĐÚNG một ngày VN — hoặc RỖNG khi không đọc được.

    Một đường đọc sổ duy nhất cho cả `_url_ca_sang` (lọc bản tối) lẫn `gop_tin_ca_toi`
    (gộp vào bản sáng). Hai nơi tự mở sổ riêng thì lệch nhau âm thầm.
    """
    try:
        sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
        from so_da_gui import url_da_gui_buoi
        return url_da_gui_buoi(buoi, ngay) or set()
    except Exception as e:                  # noqa: BLE001
        print(f"Không đọc được sổ buổi {buoi} ngày {ngay} ({e}) — không lọc gì.",
              file=sys.stderr)
        return set()


def gop_tin_ca_toi(items, cur, kind, now):
    """Bản SÁNG gộp thêm tin quét ở ca TỐI hôm qua (Huy chốt 26/08/2026).

    Nguyên văn: *"từ giờ bản tin 4h sáng hãy gộp cả tin quét được lúc 9h tối vào, nhớ đối
    chiếu với cả file Jay Lâm gửi để chống trùng lặp"*.

    Vì sao tin ca tối vắng mặt trong bản sáng: `pick_items` lấy HỢP của (mới so với commit
    cha) và (`_addedDate == generatedAt`). Sáng nay `generatedAt` là ngày MỚI nên tin nạp
    tối qua không phải "hôm nay", còn commit cha lại chính là commit của lô tối qua nên
    chúng cũng không "mới" — rơi khỏi cả hai vế, không lệnh nào báo.

    ⚠️ **Chỉ trừ tin đã gửi ở ca SÁNG HÔM QUA, không trừ theo dòng `toi`.** Trừ theo `toi`
    là xoá đúng nhóm tin vừa được lệnh gộp vào. Nhóm phải loại là bản sáng hôm qua — lặp
    lại nó nghĩa là đọc cùng một tin hai buổi sáng liền.

    ⚠️ **FAIL VỀ PHÍA GỘP DƯ:** sổ thiếu hoặc đọc hỏng thì không trừ được gì, bản sáng lặp
    lại tin của bản sáng hôm qua — Huy thấy ngay khi đọc. Hướng ngược lại là mất tin trong
    im lặng.

    Bản TỐI không gọi hàm này: tin hôm qua đã đi trong bản tin của chính hôm qua.
    """
    if la_buoi_toi(now):
        return items
    lst = event_items(cur) if kind == "events" else (cur.get(kind, []) or [])
    hom_qua = (now - datetime.timedelta(days=1)).strftime("%Y-%m-%d")
    da_co = {_khoa_tin(it) for it in items}
    ung_vien = [it for it in lst
                if _khoa_tin(it) not in da_co
                and (it.get("_addedDate") or it.get("date")) == hom_qua]
    if not ung_vien:
        return items
    # Đọc sổ CHỈ khi có ứng viên: phiên nào cũng gọi thì log đầy dòng "không đọc được sổ"
    # của những ngày vốn chẳng có gì để gộp.
    da_gui_sang = _doc_url_buoi("sang", hom_qua)
    them = [it for it in ung_vien
            if (it.get("sourceUrl") or "").strip() not in da_gui_sang]
    if not them:
        return items
    giu = da_co | {_khoa_tin(it) for it in them}
    out = [it for it in lst if _khoa_tin(it) in giu]
    print(f"Ca tối {hom_qua}: gộp thêm {len(them)} tin vào bản sáng: "
          + "; ".join((it.get("title") or it.get("sourceUrl") or "?")[:70] for it in them),
          file=sys.stderr)
    return out


def loc_bo_tin_ca_sang(items, now):
    """Bỏ khỏi bản TỐI những tin đã gửi ở ca SÁNG cùng ngày (Huy chốt 01/08/2026).

    ⚠️ ĐẢO LẠI chú thích cũ trong `main()` (*"KHÔNG lọc sổ ở đây… đừng cho nhất quán bằng
    cách bọc loc_chua_gui vào"*). Chú thích đó viết 27/07 khi THÂN EMAIL còn sống và chính
    nó gánh vai thực thi luật *"email tối = tin cả ngày TRỪ tin ca sáng sớm"*. Từ khi
    `GUI_EMAIL='0'` (27/07) thân email tắt và `.docx` thành kênh DUY NHẤT mang nội dung —
    tức lớp thực thi luật biến mất trong im lặng, còn luật thì vẫn nằm trong CLAUDE.md.
    Đo thật 01/08 trên sổ: **100% tin ca sáng lặp lại trong bản tối, cả 4/4 ngày** (28/07
    9/9 · 29/07 17/17 · 30/07 16/16 · 31/07 6/6).

    Phần chú thích cũ VẪN ĐÚNG ở chỗ nó cảnh báo: đừng bọc `loc_chua_gui` vào đây. Hàm đó
    lọc theo TOÀN sổ nên sẽ giết cả bản dựng lại của chính ca tối. Ở đây chỉ đọc dòng
    `buoi == "sang"` của ĐÚNG ngày hôm nay — xem `so_da_gui.url_da_gui_buoi`.

    Bản SÁNG không gọi hàm này: tin của nó vừa nạp, chưa từng gửi.
    """
    da_gui = _url_ca_sang(now)
    if not da_gui:
        return items
    out = [it for it in items if it.get("sourceUrl") not in da_gui]
    if len(out) != len(items):
        print(f"Ca sáng: bỏ {len(items) - len(out)} tin đã gửi trong bản tin sáng nay.",
              file=sys.stderr)
    return out


# --- File Jay Lâm gửi = BỘ LỌC: bỏ tin CỦA MÌNH mà anh ta đã có ----------
# Huy đảo nguyên tắc 01/08/2026 (xem docstring đầu file). Mục 5 và toàn bộ đường đọc Supabase
# đã bỏ; thứ còn lại ở đây là lớp MỎNG đọc sổ do phiên quét khai.
#
# ⚠️ Vì sao phép lọc không nằm ở file này: so LINK THUẦN là vô dụng — đo 01/08 trên 12 tin
# quét và 37 URL trong file Jay ra **0 tin trùng URL**, trong khi đọc hiểu ra **03 tin trùng
# sự kiện** (Mahan Air · tuần tra Scarborough · NITE-STAR 981 triệu USD), vì Jay Lâm viết lại
# bằng tiếng Việt từ nguồn khác hẳn. Đọc hiểu theo sự kiện cần agent, mà file này chạy trong
# workflow `notify-email.yml` — không có agent. Nên phiên quét khai kết quả vào sổ, file này
# chỉ so URL với sổ đó.
#
# ⚠️ Hằng số khung ngày (`JAYLAM_MAX_AGE_DAYS…`) đã XOÁ khỏi file này, cố ý: sổ tự cắt theo
# `tin_jaylam.py::GIU_NGAY`, nên ở đây không còn phép đo tuổi nào. Nếu `HeThong/dong-bo-luat.py`
# còn khai file này là nơi trích khung ngày Điểm Tin thì phải gỡ dòng đó, kẻo nó báo
# "KHÔNG ĐO ĐƯỢC" vĩnh viễn.
SO_TRUNG_JAYLAM = "logs/trung-jaylam.json"


def la_buoi_toi(now):
    """True nếu `now` (giờ VN) rơi vào buổi TỐI — cùng ngưỡng 14h với `ten_file()`.

    Tách thành hàm riêng để test được KHÔNG cần giả lập `datetime.now()`: truyền thẳng
    một mốc giờ cụ thể vào là đủ.

    ⚠️ Mục 5 đã bỏ nên hàm này CHỈ còn phục vụ `ten_file()`. Đừng thấy hết chỗ gọi trong
    logic mục 5 mà xoá — tên file `-sang-som-5h-`/`-toi-21h-` vẫn dựa vào ngưỡng này.
    """
    return now.hour >= 14


def doc_url_trung_jaylam():
    """Tập URL tin CỦA MÌNH đã được phiên quét xác định là trùng file Jay Lâm.

    Nguồn: `logs/trung-jaylam.json`, do `scripts/tin_jaylam.py --ghi-loai` ghi sau khi agent
    đối chiếu theo SỰ KIỆN. Sổ tự cắt theo `GIU_NGAY` bên đó nên ở đây không đo tuổi.

    ⚠️ FAIL VỀ PHÍA KHÔNG LỌC, và đó là chủ ý: sổ thiếu/hỏng thì bản tin LẶP một tin Jay Lâm
    đã có — phiền nhưng Huy thấy được ngay khi đọc; còn fail về phía lọc thì tin biến mất
    khỏi bản tin mà không ai biết. Xoá tin là hướng lệch tệ hơn lặp tin.
    ⚠️ Nhưng KHÔNG được im: thiếu sổ vẫn in một dòng, kẻo bước `--ghi-loai` bị bỏ nhiều phiên
    liền mà nhìn log không phân biệt được với "hôm nay không có tin nào trùng".
    """
    try:
        with open(SO_TRUNG_JAYLAM, "r", encoding="utf-8") as f:
            rows = json.load(f)
    except FileNotFoundError:
        print(f"Không có {SO_TRUNG_JAYLAM} — không lọc tin trùng file Jay Lâm. "
              "(Bước `tin_jaylam.py --ghi-loai` của phiên quét chưa từng chạy?)",
              file=sys.stderr)
        return set()
    except (OSError, json.JSONDecodeError) as e:
        print(f"Đọc {SO_TRUNG_JAYLAM} hỏng ({e}) — KHÔNG lọc, giữ nguyên tin.",
              file=sys.stderr)
        return set()
    if not isinstance(rows, list):
        print(f"{SO_TRUNG_JAYLAM} không phải mảng — KHÔNG lọc, giữ nguyên tin.",
              file=sys.stderr)
        return set()
    return {(r.get("url") or "").strip() for r in rows
            if isinstance(r, dict) and (r.get("url") or "").strip()}


def loc_bo_trung_jaylam(items, urls, ten_muc=""):
    """Bỏ khỏi bản tin những tin mà Jay Lâm đã có (Huy chốt 01/08/2026).

    So bằng `sourceUrl` — ở ĐÂY thì so URL là đúng, khác hẳn việc dùng URL để PHÁT HIỆN
    trùng. Cái được so là URL tin của chính mình, do agent liệt kê ra sau khi đã đọc hiểu;
    nó chỉ đóng vai khoá tra cứu, không phải phép đo giống nhau.

    Kêu từng tin bị bỏ: xoá tin là mất nội dung, phải soi ngược được (sổ giữ cả `trung_voi`).
    """
    if not urls:
        return items
    out = [it for it in items if (it.get("sourceUrl") or "").strip() not in urls]
    if len(out) != len(items):
        bo = [it for it in items if (it.get("sourceUrl") or "").strip() in urls]
        print(f"Bộ lọc Jay Lâm{' [' + ten_muc + ']' if ten_muc else ''}: bỏ {len(bo)} tin "
              "Jay Lâm đã có: "
              + "; ".join((it.get("title") or it.get("sourceUrl") or "?")[:70] for it in bo),
              file=sys.stderr)
    return out


def main(now=None):
    now = now or datetime.datetime.now(VN)
    with open("index.html", "r", encoding="utf-8") as f:
        cur = extract_data(f.read())
    prev = prev_data()

    # ⚠️ ĐỪNG bọc `loc_chua_gui` vào đây — nó lọc theo TOÀN sổ nên giết cả bản dựng lại của
    # chính ca tối. Thứ đúng là `loc_bo_tin_ca_sang`: chỉ bỏ tin đã gửi ở ca SÁNG CÙNG NGÀY
    # (Huy chốt 01/08/2026 sau khi bắt được tin Healio/Schwartz-CDC lặp ở cả hai bản). Chỉ
    # thị 27/07 *"file word gộp cả 11 tin hôm nay"* nói về tin quét TAY giữa ngày — nhóm đó
    # không ghi sổ nên vẫn được giữ nguyên. Xem docstring của hàm.
    us = loc_bo_tin_ca_sang(pick_items(cur, prev, "usNews"), now)
    world = loc_bo_tin_ca_sang(pick_items(cur, prev, "worldNews"), now)
    events = loc_bo_tin_ca_sang(pick_items(cur, prev, "events"), now)

    # Bản SÁNG gộp thêm tin của ca TỐI hôm qua (Huy chốt 26/08/2026). Đặt SAU
    # `loc_bo_tin_ca_sang` (hàm đó chỉ chạm bản tối) và TRƯỚC bộ lọc Jay Lâm — tin gộp thêm
    # phải đi qua bộ lọc ấy, vì file Jay Lâm thường tới SAU bản tin tối.
    us = gop_tin_ca_toi(us, cur, "usNews", now)
    world = gop_tin_ca_toi(world, cur, "worldNews", now)
    events = gop_tin_ca_toi(events, cur, "events", now)

    # Bộ lọc Jay Lâm — bỏ tin mình quét được mà anh ta đã có (Huy chốt 01/08/2026).
    # Áp cho CẢ BA mục: cơ chế gây vấp của `loc_bo_tin_ca_sang` bản đầu là chỉ áp 03 mục quét
    # mà quên mục 5, nên lỗ trùng vẫn hở; ở đây không còn mục 5 nhưng vẫn phải phủ đủ ba,
    # tin Jay Lâm gửi có cả tin thế giới lẫn tin tập trận.
    trung_jl = doc_url_trung_jaylam()
    us = loc_bo_trung_jaylam(us, trung_jl, "usNews")
    world = loc_bo_trung_jaylam(world, trung_jl, "worldNews")
    events = loc_bo_trung_jaylam(events, trung_jl, "events")

    sections = build_sections(us, world, events)

    total = sum(len(items) for _, items in sections)
    if total == 0:
        print("DOCX=")
        return

    gen = cur.get("generatedAt", "")

    doc = Document()
    # Khổ A4 + lề 1,0" cả bốn phía — đo trên mẫu 01/09/2026 (pgSz 11909x16834 twips,
    # pgMar 1440 cả bốn). python-docx mặc định khổ Letter, không đặt là bản in lệch.
    for s in doc.sections:
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)

    # Style Normal: Times New Roman 14pt — mẫu khai ở style chứ không chỉ ở từng run, nên
    # đoạn nào lỡ quên `set_font` vẫn ra đúng chữ thay vì rơi về Calibri 11.
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = Pt(SIZE)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # ⛔ KHÔNG in dòng tiêu đề "ĐIỂM TIN NGÀY d.M.yyyy" — mẫu 01/09/2026 vào THẲNG mục (1).
    # Ngày của bản tin nằm ở TÊN FILE (`ten_file`), nên bỏ dòng này không mất thông tin nào.
    # Thêm lại là lệch mẫu ngay dòng đầu tiên, chỗ người duyệt nhìn trước nhất.
    idx = 0
    for name, items in sections:
        if not items:
            continue
        idx += 1
        add_dau_muc(doc, f"({idx}) {name}")
        if name == MUC_DIA_BAN:
            # Mục địa bàn in theo 03 tiểu mục; tiểu mục rỗng thì bỏ hẳn nhãn, không in
            # nhãn trống (mẫu không có nhãn nào đứng một mình).
            for tm in THU_TU_TIEU_MUC:
                nhom = [it for it in items if it.get(KHOA_TIEU_MUC) == tm]
                if not nhom:
                    continue
                add_tieu_muc(doc, tm)
                for it in nhom:
                    add_item(doc, it)
        else:
            for it in items:
                add_item(doc, it)
        add_dong_trong(doc)

    out = f"/tmp/{ten_file(gen, now)}"
    doc.save(out)
    # ⚠️ KHÔNG còn bước đánh dấu Supabase nào ở đây (bỏ 01/08/2026 cùng mục 5). Việc đóng sổ
    # dòng Jay Lâm hết khung ngày nay nằm ở `scripts/tin_jaylam.py::in_hang_cho()` — tức ở
    # chính chỗ đọc hàng chờ, thay vì ở chỗ dựng file. Đừng dựng lại đường ghi Supabase từ
    # file này: nó chạy trong workflow `notify-email.yml` mà workflow đó có thể chạy nhiều
    # lần cho cùng một bản tin.
    print(f"DOCX={out}")


if __name__ == "__main__":
    main()
