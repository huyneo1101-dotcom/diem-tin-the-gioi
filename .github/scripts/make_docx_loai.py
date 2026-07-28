#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Tạo file .docx THỨ HAI gửi kèm bản tin: **TIN BỊ LOẠI dù thuộc đúng 5 chủ đề**, kèm LÝ DO.

Chỉ thị Huy 28/07/2026: *"từ giờ mỗi khi gửi hãy gửi thêm 1 file word nữa, trong đó gồm
các tin đã bị loại dù thuộc đúng 5 chủ đề. ghi rõ lý do bị loại."*

VÌ SAO CẦN: bản tin chỉ cho thấy thứ ĐƯỢC nhận. Phần bị loại — tin đúng gu nhưng rớt vì
ngoài khung ngày, nghi trùng, nguồn không khớp — trước nay chỉ nằm trong `logs/loai-tin.md`
trên GitHub, Huy phải tự vào repo đọc. Không ai rà thì loại nhầm sẽ không bao giờ bị bắt.

━━━ HAI NGUỒN DỮ LIỆU, ƯU TIÊN THEO THỨ TỰ ━━━
1. `logs/loai-tin.json`  ← NGUỒN CHÍNH, có cấu trúc, do phiên quét ghi (cùng lúc với
   `logs/scan-gaps.json`). Mỗi tin một mục: chủ đề · tiêu đề · nguồn · link · ngày · lý do.
2. `logs/loai-tin.md`    ← FALLBACK. Trích NGUYÊN VĂN các gạch đầu dòng dưới mục ngày
   tương ứng. Dùng khi phiên quét chưa kịp ghi JSON (phiên cũ, hoặc agent quên).

⚠️ VÌ SAO PHẢI CÓ FALLBACK, đừng "dọn cho gọn" bỏ đi: nguồn chính phụ thuộc agent ghi
đúng định dạng. Nếu chỉ đọc JSON thì hôm nào agent quên là Huy nhận file RỖNG — mà rỗng
trông y hệt "hôm nay không loại tin nào", tức lỗi câm. `loai-tin.md` thì agent nào cũng
ghi (đã thành thói quen từ 13/07), nên nó là lưới an toàn.

⚠️ CHỐT AN TOÀN NGÀY — giống `scan-gaps.json` trong `send-email.js`: `date` trong JSON
phải KHỚP `DATA.generatedAt`, lệch thì BỎ nguồn 1 và rơi xuống fallback. Không có chốt này
thì hôm nào phiên quét không ghi file mới, Huy sẽ nhận lại nguyên lý do của HÔM TRƯỚC mà
tưởng là của hôm nay.

Xuất: dòng cuối stdout `DOCX_LOAI=<path>`; không có gì để ghi -> `DOCX_LOAI=` rỗng
(im lặng ĐÚNG, không phải lỗi — có ngày không loại tin nào).

Chạy: python3 .github/scripts/make_docx_loai.py
"""
import json
import pathlib
import re
import sys

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
import make_docx as md  # noqa: E402  — dùng chung ten_file/set_font/add_hyperlink/extract_data

ROOT = pathlib.Path(__file__).resolve().parents[2]

# Thứ tự mục trong file — BÁM ĐÚNG 5 chủ đề của phạm vi quét (CLAUDE.md banner 23/07).
# Tin không khớp chủ đề nào rơi vào "Khác" ở cuối chứ KHÔNG bị nuốt: một tin bị loại mà
# biến mất khỏi cả file tin-bị-loại thì đúng là thứ file này sinh ra để chống.
THU_TU = ["Nội bộ Mỹ", "Úc & Biển Đông", "CNQS Mỹ", "Mỹ – Mali", "Predator's Run 2026"]

# Tên chủ đề agent hay viết tắt/viết khác -> tên chuẩn ở trên.
DONG_NGHIA = {
    "noi bo my": "Nội bộ Mỹ",
    "nội bộ mỹ": "Nội bộ Mỹ",
    "us": "Nội bộ Mỹ",
    "uc & bien dong": "Úc & Biển Đông",
    "úc & biển đông": "Úc & Biển Đông",
    "úc và biển đông": "Úc & Biển Đông",
    "uc va bien dong": "Úc & Biển Đông",
    "biển đông": "Úc & Biển Đông",
    "cnqs my": "CNQS Mỹ",
    "cnqs mỹ": "CNQS Mỹ",
    "cnqs": "CNQS Mỹ",
    "qs-khcn": "CNQS Mỹ",
    "my - mali": "Mỹ – Mali",
    "mỹ - mali": "Mỹ – Mali",
    "mỹ – mali": "Mỹ – Mali",
    "my – mali": "Mỹ – Mali",
    "mali": "Mỹ – Mali",
    "predator's run": "Predator's Run 2026",
    "predator's run 2026": "Predator's Run 2026",
    "predators run": "Predator's Run 2026",
    "predator": "Predator's Run 2026",
}


def chuan_chu_de(s):
    """Đưa tên chủ đề agent ghi về một trong 5 tên chuẩn; không nhận ra thì giữ nguyên."""
    k = (s or "").strip()
    return DONG_NGHIA.get(k.lower(), k or "Khác")


def doc_json(generated_at):
    """Nguồn CHÍNH. Trả `(items, dung_duoc)`.

    ⚠️ HAI GIÁ TRỊ chứ không phải một danh sách — vì `[]` gộp mất hai ca KHÁC HẲN nhau
    (vấp thật lúc nghiệm thu 28/07):
      • file hợp lệ, `items: []`  = hôm nay KHÔNG loại tin nào -> im lặng ĐÚNG, **đừng**
        rơi xuống fallback (rơi xuống là dựng lại file từ `loai-tin.md`, tức moi ra đúng
        thứ phiên quét đã cố ý khai là không có);
      • thiếu file / JSON hỏng / lệch ngày = KHÔNG BIẾT hôm nay loại gì -> phải fallback.
    Cùng lớp lỗi với nhánh `.docx` của `send_telegram.py` ("0 tin" ≠ "dựng hỏng") và với
    chốt secret Telegram ("chưa cấu hình" ≠ "cấu hình gãy").
    """
    p = ROOT / "logs" / "loai-tin.json"
    try:
        d = json.loads(p.read_text(encoding="utf-8"))
    except FileNotFoundError:
        print("[loai] chưa có logs/loai-tin.json — dùng fallback loai-tin.md", file=sys.stderr)
        return [], False
    except Exception as e:  # noqa: BLE001
        print(f"[loai] logs/loai-tin.json lỗi ({e}) — dùng fallback loai-tin.md", file=sys.stderr)
        return [], False
    ngay = str(d.get("date", "")).strip()
    if generated_at and ngay != generated_at:
        print(f"[loai] loai-tin.json ngày {ngay!r} ≠ generatedAt {generated_at!r} — BỎ, "
              "tránh gửi lý do của hôm trước. Dùng fallback.", file=sys.stderr)
        return [], False
    items = d.get("items")
    if not isinstance(items, list):
        print("[loai] loai-tin.json: 'items' không phải mảng — bỏ, dùng fallback.", file=sys.stderr)
        return [], False
    return ([x for x in items if isinstance(x, dict) and (x.get("title") or x.get("reason"))],
            True)


def doc_markdown(generated_at, buoi_toi):
    """FALLBACK: trích nguyên văn gạch đầu dòng dưới mục ngày trong logs/loai-tin.md.

    Mục có dạng `## 2026-07-28 (phiên TỐI — CI 21:00)`. Cùng một ngày có thể có CẢ mục
    phiên sáng lẫn phiên tối, nên chọn theo buổi đang chạy; không có mục khớp buổi thì
    lấy mục CUỐI CÙNG của ngày đó (thà đưa nội dung gần đúng còn hơn trả rỗng).
    """
    p = ROOT / "logs" / "loai-tin.md"
    try:
        raw = p.read_text(encoding="utf-8")
    except Exception as e:  # noqa: BLE001
        print(f"[loai] không đọc được loai-tin.md ({e}).", file=sys.stderr)
        return []

    khoi = []  # [(tieu_de_muc, [dòng...])]
    hien = None
    for ln in raw.splitlines():
        if ln.startswith("## "):
            tieu = ln[3:].strip()
            if generated_at and tieu.startswith(generated_at):
                hien = (tieu, [])
                khoi.append(hien)
            else:
                hien = None
            continue
        if hien is not None and ln.strip().startswith("- "):
            hien[1].append(ln.strip()[2:].strip())

    if not khoi:
        return []
    tu = "tối" if buoi_toi else "sáng"
    chon = None
    for tieu, dong in khoi:
        if tu in tieu.lower():
            chon = (tieu, dong)
    if chon is None:
        chon = khoi[-1]
    print(f"[loai] fallback: lấy mục {chon[0]!r} trong loai-tin.md ({len(chon[1])} dòng).",
          file=sys.stderr)
    return [{"_raw": d} for d in chon[1]]


def them_muc(doc, ten, items):
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ph.paragraph_format.space_before = Pt(8)
    ph.paragraph_format.space_after = Pt(4)
    md.set_font(ph.add_run(ten), size=md.SIZE, bold=True)

    for it in items:
        if it.get("_raw"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(8)
            md.set_font(p.add_run("- " + it["_raw"]), size=md.SIZE)
            continue

        tit = (it.get("title") or "").strip()
        ngn = (it.get("sourceName") or "").strip()
        ngay = md.ngay_ngan(it.get("date"))
        dau = " · ".join(x for x in (ngn, ngay) if x)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(2)
        md.set_font(p.add_run(f"- {tit}" + (f" ({dau})" if dau else "")), size=md.SIZE)

        ly_do = (it.get("reason") or "").strip()
        if ly_do:
            pl = doc.add_paragraph()
            pl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pl.paragraph_format.space_after = Pt(2)
            md.set_font(pl.add_run("Lý do loại: "), size=md.SIZE, bold=True)
            md.set_font(pl.add_run(ly_do), size=md.SIZE, italic=True)

        url = (it.get("sourceUrl") or "").strip()
        if url:
            pu = doc.add_paragraph()
            pu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pu.paragraph_format.space_after = Pt(8)
            md.add_hyperlink(pu, url, url)
        else:
            p.paragraph_format.space_after = Pt(8)


def main():
    try:
        cur = md.extract_data((ROOT / "index.html").read_text(encoding="utf-8"))
    except Exception as e:  # noqa: BLE001
        print(f"❌ [loai] không đọc được DATA trong index.html: {e}", file=sys.stderr)
        return 1
    gen = cur.get("generatedAt", "")

    import datetime
    buoi_toi = datetime.datetime.now(md.VN).hour >= 14

    items, dung_duoc = doc_json(gen)
    nguon = "logs/loai-tin.json"
    if dung_duoc and not items:
        print("[loai] loai-tin.json khai RÕ hôm nay không loại tin nào (items rỗng) — "
              "không dựng file, KHÔNG fallback.", file=sys.stderr)
        print("DOCX_LOAI=")
        return 0
    if not dung_duoc:
        items = doc_markdown(gen, buoi_toi)
        nguon = "logs/loai-tin.md (fallback)"

    if not items:
        print("[loai] không có tin bị loại nào cho ngày này — không dựng file.", file=sys.stderr)
        print("DOCX_LOAI=")
        return 0

    # Gom theo chủ đề, giữ thứ tự 5 chủ đề; phần không nhận ra chủ đề xuống cuối.
    nhom = {}
    for it in items:
        ten = "Ghi chú phiên quét" if it.get("_raw") else chuan_chu_de(it.get("chu_de"))
        nhom.setdefault(ten, []).append(it)
    thu_tu = [t for t in THU_TU if t in nhom] + [t for t in nhom if t not in THU_TU]

    try:
        y, m, d = gen.split("-")
        title_date = f"{int(d)}.{int(m)}.{y}"
    except Exception:  # noqa: BLE001
        title_date = gen

    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(1.25)
        s.right_margin = Inches(1.25)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)

    pt = doc.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_after = Pt(4)
    md.set_font(pt.add_run(f"TIN BỊ LOẠI — NGÀY {title_date}"), size=md.SIZE, bold=True)

    pg = doc.add_paragraph()
    pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pg.paragraph_format.space_after = Pt(10)
    md.set_font(pg.add_run("Tin thuộc đúng 5 chủ đề nhưng không được đưa vào bản tin, kèm lý do"),
                size=md.SIZE, italic=True)

    for ten in thu_tu:
        them_muc(doc, ten, nhom[ten])

    out = f"/tmp/{md.ten_file(gen, loai=True)}"
    doc.save(out)
    print(f"[loai] nguồn: {nguon} — {len(items)} mục, {len(thu_tu)} nhóm.", file=sys.stderr)
    print(f"DOCX_LOAI={out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
