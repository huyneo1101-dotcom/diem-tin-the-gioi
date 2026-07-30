#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Test cho tính năng "Jay Lâm gửi tin qua bot -> gộp cuối ngày" (thêm 30/07/2026).

    python3 tests/test-gop-tin-jaylam.py

Không chạm mạng thật: mọi hàm gọi curl/Supabase/Telegram đều bị monkeypatch. Canh 3 mảnh:
`scripts/docx_text.py` (bóc chữ từ .docx) · `scripts/gop_tin_jaylam.py::dung_docx/main` (gộp
cuối ngày) · `scripts/telegram_bot.py::xu_ly_tin_jaylam` (nhận file từ bot).

Theo mục 17 CLAUDE.md toàn cục: các ca 6-10 là ca PHẢI CHẶN — chứng minh script KHÔNG gửi gì
(và KHÔNG lưu gì) khi gặp điều kiện xấu, không chỉ chứng minh nó chạy trơn ở luồng bình thường.
"""
import os
import pathlib
import sys
import tempfile
import unittest.mock as mock

ROOT = pathlib.Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT / "scripts"))

import docx_text  # noqa: E402
import gop_tin_jaylam as g  # noqa: E402
import telegram_bot as tb  # noqa: E402
from docx import Document  # noqa: E402

CA = []


def kiem(ten, dat):
    dat = bool(dat)
    CA.append((ten, dat))
    print(("✓" if dat else "✗") + " " + ten)


def docx_mau(doan):
    d = Document()
    for p in doan:
        d.add_paragraph(p)
    tmp = tempfile.mktemp(suffix=".docx")
    d.save(tmp)
    return tmp


# --- docx_text.trich() ---------------------------------------------------
kiem("trich() file không tồn tại -> rỗng, không crash",
     docx_text.trich("/khong/co/that.docx") == "")

f1 = docx_mau(["Dòng một có dấu: Đà Nẵng", "Dòng hai"])
t1 = docx_text.trich(f1)
kiem("trich() giữ dấu tiếng Việt", "Đà Nẵng" in t1)
kiem("trich() giữ đủ các dòng", "Dòng một" in t1 and "Dòng hai" in t1)

f2 = docx_mau(["a" * 100])
t2 = docx_text.trich(f2, max_chars=10)
kiem("trich() cắt đúng độ dài + có dấu …", len(t2) == 11 and t2.endswith("…"))

f3 = tempfile.mktemp(suffix=".docx")
pathlib.Path(f3).write_text("không phải file zip/docx thật", encoding="utf-8")
kiem("trich() file .docx giả (không phải zip) -> rỗng", docx_text.trich(f3) == "")

# --- gop_tin_jaylam.dung_docx() ------------------------------------------
rows = [{"id": 1, "ten_file": "a.docx", "ten": "Jay Lâm", "noi_dung": "Nội dung A",
         "created_at": "2026-07-30T03:00:00Z"}]
path = g.dung_docx(rows, "2026-07-30")
full = "\n".join(p.text for p in Document(path).paragraphs)
kiem("dung_docx() có tiêu đề ngày", "2026-07-30" in full)
kiem("dung_docx() có tên file gốc", "a.docx" in full)
kiem("dung_docx() có nội dung", "Nội dung A" in full)

# --- gop_tin_jaylam.main() — CA PHẢI CHẶN --------------------------------
ENV_DU = {"TELEGRAM_BOT_TOKEN": "TOK", "TELEGRAM_CHAT_ID": "111,222"}

with mock.patch.dict(os.environ, ENV_DU, clear=False), \
        mock.patch.object(g, "doc_ngay", return_value=[]), \
        mock.patch.object(g, "send_document") as sd, \
        mock.patch.object(tb, "_anon_key", return_value="k"), \
        mock.patch.object(tb, "_dt_bot_key", return_value="dk"):
    ma = g.main()
kiem("main() không có tin trong ngày -> mã 10, KHÔNG gửi", ma == 10 and not sd.called)

with mock.patch.dict(os.environ, ENV_DU, clear=False), \
        mock.patch.object(g, "doc_ngay", return_value=None), \
        mock.patch.object(g, "send_document") as sd, \
        mock.patch.object(tb, "_anon_key", return_value="k"), \
        mock.patch.object(tb, "_dt_bot_key", return_value="dk"):
    ma = g.main()
kiem("main() đọc Supabase hỏng -> mã 1, KHÔNG gửi", ma == 1 and not sd.called)

with mock.patch.dict(os.environ, ENV_DU, clear=False), \
        mock.patch.object(g, "doc_ngay", return_value=rows), \
        mock.patch.object(g, "send_document") as sd, \
        mock.patch.object(tb, "_anon_key", return_value="k"), \
        mock.patch.object(tb, "_dt_bot_key", return_value="dk"), \
        mock.patch.object(tb, "chat_chu", return_value=""):
    ma = g.main()
kiem("main() thiếu chat_chu() (Huy) -> mã 1, KHÔNG gửi", ma == 1 and not sd.called)

with mock.patch.dict(os.environ, {"TELEGRAM_BOT_TOKEN": "", "TELEGRAM_CHAT_ID": ""},
                      clear=False), \
        mock.patch.object(g, "doc_ngay", return_value=rows), \
        mock.patch.object(g, "send_document") as sd:
    ma = g.main()
kiem("main() thiếu secret Telegram -> mã 1 (job ĐỎ), KHÔNG gửi", ma == 1 and not sd.called)

with mock.patch.dict(os.environ, ENV_DU, clear=False), \
        mock.patch.object(g, "doc_ngay", return_value=rows), \
        mock.patch.object(g, "send_document",
                           return_value={"ok": True}) as sd, \
        mock.patch.object(g, "danh_dau_da_gop") as dau, \
        mock.patch.object(tb, "_anon_key", return_value="k"), \
        mock.patch.object(tb, "_dt_bot_key", return_value="dk"), \
        mock.patch.object(tb, "chat_chu", return_value="111"):
    ma = g.main()
kiem("main() luồng bình thường -> mã 0, gửi đúng 1 lần, đánh dấu đã gộp",
     ma == 0 and sd.call_count == 1 and sd.call_args.args[1] == "111" and dau.called)

# --- telegram_bot.xu_ly_tin_jaylam() — CA PHẢI CHẶN ----------------------
with mock.patch.object(tb, "call") as call_m, \
        mock.patch.object(tb, "tai_file") as tai_m, \
        mock.patch.object(tb, "luu_tin_jaylam") as luu_m:
    tb.xu_ly_tin_jaylam("111", "111", {"from": {"first_name": "Jay"}},
                         {"file_name": "anh.jpg", "file_id": "abc"})
kiem("xu_ly_tin_jaylam() từ chối file không phải .docx -> KHÔNG tải, KHÔNG lưu",
     not tai_m.called and not luu_m.called and call_m.called
     and ".docx" in call_m.call_args.args[2]["text"])

with mock.patch.object(tb, "call") as call_m, \
        mock.patch.object(tb, "tai_file", return_value=False) as tai_m, \
        mock.patch.object(tb, "luu_tin_jaylam") as luu_m:
    tb.xu_ly_tin_jaylam("111", "111", {"from": {"first_name": "Jay"}},
                         {"file_name": "tin.docx", "file_id": "abc"})
kiem("xu_ly_tin_jaylam() tải file hỏng -> KHÔNG lưu, có báo lỗi",
     tai_m.called and not luu_m.called and "hỏng" in call_m.call_args.args[2]["text"])

f4 = docx_mau(["Tin thật từ Jay Lâm"])
with mock.patch.object(tb, "call") as call_m, \
        mock.patch.object(tb, "tai_file",
                           side_effect=lambda tok, fid, dich: pathlib.Path(f4).rename(dich)
                           or True) as tai_m, \
        mock.patch.object(tb, "luu_tin_jaylam", return_value=True) as luu_m:
    tb.xu_ly_tin_jaylam("111", "111", {"from": {"first_name": "Jay Lâm"}},
                         {"file_name": "tin.docx", "file_id": "abc"})
kiem("xu_ly_tin_jaylam() luồng bình thường -> lưu đúng nội dung, báo đã nhận",
     luu_m.called and "Jay Lâm" in luu_m.call_args.args[1]
     and "Tin thật" in luu_m.call_args.args[3]
     and "Đã nhận" in call_m.call_args.args[2]["text"])

so_dat = sum(1 for _, ok in CA if ok)
print(f"\n{so_dat}/{len(CA)} ca đạt")
sys.exit(0 if so_dat == len(CA) else 1)
