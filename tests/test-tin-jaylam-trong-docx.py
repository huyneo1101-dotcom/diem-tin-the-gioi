#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Test cho mục 5 "Tin Jay Lâm gửi" trong file .docx bản tin.

    python3 tests/test-tin-jaylam-trong-docx.py
    python3 tests/test-tin-jaylam-trong-docx.py --tu-kiem   # chứng minh test BẮT ĐƯỢC lỗi

Thiết kế mục 5 đã đổi 30/07/2026 theo bảng chọn Huy chốt — bộ test canh đúng 4 quyết định đó:
  (1) tin Jay Lâm in theo KHUÔN TIN CHUẨN (tóm tắt + link), do phiên quét viết qua
      `scripts/tin_jaylam.py`; dòng CHƯA qua bước đó thì KHÔNG vào file, nằm chờ và không
      đóng sổ (Huy chốt 01/08/2026 — bỏ hẳn nhánh dán nguyên văn, xem `tach_chua_tom_tat`);
  (2) có DÒNG NHÃN "chưa qua thang xác minh nguồn" ở đầu mục;
  (3) KHÔNG trần số lượng;
  (4) áp KHUNG 2 NGÀY như tin quét — quá hạn thì bỏ khỏi file nhưng VẪN đánh dấu `da_gop`.

Cộng 3 việc dễ hỏng câm của kiểu tích hợp này (mục 17 CLAUDE.md — hỏng thì im lặng cho qua):
mục Jay Lâm CHỈ vào bản buổi TỐI · bộ lọc chống trùng phải chặn tin quét thường ĐÃ có · tin bị
lọc trùng vẫn phải được đánh dấu `da_gop` kẻo nằm lại vĩnh viễn.

Không chạm mạng thật: `subprocess.run` của module bị thay khi đo `doc_tin_jaylam_chua_gop`,
còn các ca hiển thị thì monkeypatch thẳng hàm đó. Yêu cầu `pip3 install python-docx`.
"""
import contextlib
import datetime
import hashlib
import io
import json
import os
import pathlib
import shutil
import subprocess as SP
import sys
import tempfile
import zoneinfo

HERE = pathlib.Path(__file__).resolve().parent
REPO = HERE.parent
GS = pathlib.Path(os.environ.get("MAKEDOCX_DIR") or (REPO / ".github" / "scripts"))
sys.path.insert(0, str(GS))

import make_docx as MD          # noqa: E402
from docx import Document       # noqa: E402

VN = zoneinfo.ZoneInfo("Asia/Ho_Chi_Minh")
MD.prev_data = lambda: None     # xem test-so-da-gui.py — bỏ phụ thuộc lịch sử git thật

SANG = datetime.datetime(2026, 7, 30, 6, 0, tzinfo=VN)
TOI = datetime.datetime(2026, 7, 30, 21, 0, tzinfo=VN)

DATA_GIA = {
    "generatedAt": "2026-07-30",
    "worldNews": [{"date": "2026-07-30", "_addedDate": "2026-07-30",
                   "title": "Tin thế giới quét được", "category": "Chính trị",
                   "sourceUrl": "https://reuters.com/tin-the-gioi"}],
    "usNews": [], "exercises": [],
}

CA = []


def kiem(ten, dat):
    dat = bool(dat)
    CA.append((ten, dat))
    print(("✓" if dat else "✗") + " " + ten)


class ThuMucGia:
    """Ghim cwd vào thư mục tạm có index.html giả, dọn sạch khi thoát."""

    def __enter__(self):
        self.cu = os.getcwd()
        self.d = pathlib.Path(tempfile.mkdtemp(prefix="jaylam-docx-"))
        (self.d / "index.html").write_text(
            "<html><script>var DATA = " + json.dumps(DATA_GIA, ensure_ascii=False)
            + ";</script></html>", encoding="utf-8")
        os.chdir(self.d)
        return self.d

    def __exit__(self, *a):
        os.chdir(self.cu)
        shutil.rmtree(self.d, ignore_errors=True)
        return False


def chay(now, doc_fn, danhdau_fn, data=None):
    """Chạy MD.main() với mục 5 bị thay. `doc_fn(now)` phải trả (trong_khung, qua_han)."""
    goc_doc, goc_danhdau = MD.doc_tin_jaylam_chua_gop, MD.danh_dau_da_gop_jaylam
    MD.doc_tin_jaylam_chua_gop, MD.danh_dau_da_gop_jaylam = doc_fn, danhdau_fn
    goc_data = globals()["DATA_GIA"]
    if data is not None:
        globals()["DATA_GIA"] = data
    buf_out, buf_err = io.StringIO(), io.StringIO()
    try:
        with ThuMucGia(), contextlib.redirect_stdout(buf_out), \
                contextlib.redirect_stderr(buf_err):
            MD.main(now=now)
    finally:
        MD.doc_tin_jaylam_chua_gop, MD.danh_dau_da_gop_jaylam = goc_doc, goc_danhdau
        globals()["DATA_GIA"] = goc_data
    out, err = buf_out.getvalue(), buf_err.getvalue()
    dong = [l for l in out.splitlines() if l.startswith("DOCX=")]
    path = dong[-1][len("DOCX="):].strip() if dong else ""
    full = "\n".join(p.text for p in Document(path).paragraphs) if path else ""
    return full, err


# Neo tham chiếu bản THẬT ngay lúc nạp module. `chay()` thay `MD.doc_tin_jaylam_chua_gop`
# bằng hàm giả, nên ca nào gọi `doc_that` TỪ BÊN TRONG hàm giả đó mà đọc qua `MD.` sẽ tự gọi
# lại chính mình -> RecursionError. Đã vấp thật khi thêm ca [44].
_DOC_GOC = MD.doc_tin_jaylam_chua_gop


def doc_that(rows, now):
    """Gọi HÀM THẬT `doc_tin_jaylam_chua_gop` với thân trả về giả — nhờ vậy phép lọc khung
    ngày và câu select đều là mã production, không phải bản mô phỏng trong test."""
    class P:
        returncode = 0
        stdout = json.dumps(rows, ensure_ascii=False)
        stderr = ""

    goc = MD.subprocess.run
    goc_env = {k: os.environ.get(k) for k in ("SUPABASE_ANON_KEY", "DT_BOT_KEY")}
    MD.subprocess.run = lambda *a, **k: P()
    os.environ["SUPABASE_ANON_KEY"] = "sb_publishable_test"
    os.environ["DT_BOT_KEY"] = "ma-test"
    buf = io.StringIO()
    try:
        with contextlib.redirect_stderr(buf):
            kq = _DOC_GOC(now)
    finally:
        MD.subprocess.run = goc
        for k, v in goc_env.items():
            if v is None:
                os.environ.pop(k, None)
            else:
                os.environ[k] = v
    return kq, buf.getvalue()


def dong_jaylam(id_, noi_dung, created_at, **kw):
    r = {"id": id_, "ten": "Jay Lâm", "ten_file": f"f{id_}.docx",
         "noi_dung": noi_dung, "created_at": created_at, "da_xu_ly": False}
    r.update(kw)
    return r


def dong_xong(id_, tieu_de, created_at, **kw):
    """Dòng ĐÃ qua bước `tin_jaylam.py` — dạng duy nhất còn được vào file .docx từ
    01/08/2026. `noi_dung` giữ một chuỗi khác hẳn tiêu đề để ca nào vô tình in nguyên văn
    trở lại là lộ ra ngay."""
    kw.setdefault("tom_tat", f"Tóm tắt tin: {tieu_de}.")
    kw.setdefault("nguon_ten", "Reuters")
    kw.setdefault("nguon_url", f"https://reuters.com/bai-{id_}")
    return dong_jaylam(id_, f"NGUYÊN VĂN dòng {id_} không được lọt vào bản tin",
                       created_at, da_xu_ly=True, tieu_de=tieu_de, **kw)


# ---------------------------------------------------------------- ngưỡng buổi
kiem("[01] la_buoi_toi() sáng sớm -> False",
     MD.la_buoi_toi(datetime.datetime(2026, 7, 30, 5, 0, tzinfo=VN)) is False)
kiem("[02] la_buoi_toi() đúng ngưỡng 14h -> True",
     MD.la_buoi_toi(datetime.datetime(2026, 7, 30, 14, 0, tzinfo=VN)) is True)
kiem("[03] la_buoi_toi() tối -> True",
     MD.la_buoi_toi(datetime.datetime(2026, 7, 30, 21, 0, tzinfo=VN)) is True)

# --------------------------------------- buổi SÁNG cũng gộp (mở 30/07/2026)
# ⚠️ ĐẢO LẠI hành vi cũ. Ca 04/05 trước đây khẳng định "buổi sáng KHÔNG đụng Supabase Jay
# Lâm" — đúng với thiết kế cũ, nhưng chính thiết kế đó là lỗ: phiên quét tối chạy
# 20:47-21:26 mà Jay Lâm gửi file lúc 21:34, muộn hơn cả bản .docx cuối cùng, nên tin phải
# chờ tới 20:47 HÔM SAU — lúc đó khung ngày đã đẩy nó sang nhóm quá hạn rồi đóng sổ. Huy
# chốt 30/07/2026: "Jay Lâm gửi tin muộn sau đợt quét buổi tối thì tự động gộp tin vào bản
# tin sáng". Hai ca này nay canh chiều NGƯỢC LẠI — sửa theo hành vi mới, KHÔNG gỡ ca.
goi = {"doc": 0}
danhdau_sang = []


def doc_co_tin(now=None):
    goi["doc"] += 1
    return [dong_xong(1, "Một tin lạ Jay Lâm gửi", "2026-07-30T02:00:00Z")], []


full, _ = chay(SANG, doc_co_tin, danhdau_sang.extend)
kiem("[04] buổi sáng: VẪN gọi doc_tin_jaylam_chua_gop()", goi["doc"] == 1)
kiem("[05] buổi sáng: file CÓ mục 'Tin Jay Lâm gửi'", "Tin Jay Lâm gửi" in full)

# [44] HỒI QUY đúng kịch bản Huy nêu: file gửi 21:34 tối qua (sau khi phiên tối đã dựng xong
# bản .docx cuối) phải lên bản tin SÁNG hôm sau, không nằm chờ tới tối hôm sau rồi quá hạn.
SANG_31 = datetime.datetime(2026, 7, 31, 3, 47, tzinfo=VN)
DATA_31 = dict(DATA_GIA, generatedAt="2026-07-31",
               worldNews=[dict(DATA_GIA["worldNews"][0], date="2026-07-31",
                               _addedDate="2026-07-31")])


def doc_gui_muon(now=None):
    # 21:34 giờ VN ngày 30/07 = 14:34Z cùng ngày.
    return doc_that([dong_jaylam(9, "Tin Jay gửi lúc 21:34", "2026-07-30T14:34:00Z",
                                 da_xu_ly=True, tieu_de="Tin gửi muộn",
                                 tom_tat="Tóm tắt tin gửi muộn tối qua.",
                                 nguon_ten="Reuters",
                                 nguon_url="https://reuters.com/tin-gui-muon")], now)[0]


full44, _ = chay(SANG_31, doc_gui_muon, [].extend, data=DATA_31)
kiem("[44] tin gửi 21:34 tối qua -> lên bản tin SÁNG hôm sau (không nằm chờ tới quá hạn)",
     "Tin Jay Lâm gửi" in full44 and "Tóm tắt tin gửi muộn tối qua." in full44)

# [45] ĐỐI CHỨNG chống nới tay: mở cho bản sáng KHÔNG được kéo theo việc nới khung ngày —
# tin quá khung ở bản sáng vẫn phải bị loại y như ở bản tối.
def doc_qua_han_sang(now=None):
    return doc_that([dong_jaylam(10, "Tin quá cũ", "2026-07-27T02:00:00Z",
                                 da_xu_ly=True, tieu_de="Tin quá cũ",
                                 tom_tat="Tóm tắt tin đã quá khung ngày rồi.",
                                 nguon_ten="Reuters",
                                 nguon_url="https://reuters.com/qua-cu")], now)[0]


qh_sang = []
full45, err45 = chay(SANG_31, doc_qua_han_sang, qh_sang.extend, data=DATA_31)
kiem("[45] buổi sáng: tin quá khung ngày VẪN bị loại + vẫn đánh dấu đã gộp",
     "Tóm tắt tin đã quá khung ngày rồi." not in full45 and qh_sang == [10])

# --------------------------------- luồng bình thường: tin ĐÃ được xử lý
danhdau_goi = []


def danhdau_ghi(ids):
    danhdau_goi.extend(ids)


def doc_da_xu_ly(now=None):
    return [dong_jaylam(
        2, "Nguyên văn rất dài mà không ai muốn đọc trong bản tin",
        "2026-07-30T09:15:00Z", da_xu_ly=True,
        tieu_de="Hạ viện Mỹ thông qua dự luật ngân sách quốc phòng",
        tom_tat="Hạ viện Mỹ thông qua dự luật ngân sách quốc phòng năm tài khóa 2027 "
                "với tỷ lệ 218-210, trong đó có khoản chi cho AUKUS.",
        nguon_ten="Reuters", nguon_url="https://reuters.com/bai-cu-the-123")], []


full2, err2 = chay(TOI, doc_da_xu_ly, danhdau_ghi)
kiem("[06] tối: có mục 'Tin Jay Lâm gửi'", "Tin Jay Lâm gửi" in full2)
kiem("[07] tin ĐÃ xử lý: in TÓM TẮT, không in nguyên văn",
     "tỷ lệ 218-210" in full2 and "không ai muốn đọc" not in full2)
kiem("[08] tin ĐÃ xử lý: có URL nguồn trong file",
     "https://reuters.com/bai-cu-the-123" in full2)
kiem("[09] tin ĐÃ xử lý: KHÔNG in cảnh báo chưa-tóm-tắt",
     "CHƯA được phiên quét" not in err2)
kiem("[10] nhãn xác minh có mặt trong file", MD.JAYLAM_NHAN_XAC_MINH in full2)
kiem("[11] nhãn thời gian ghi cả NGÀY, không chỉ giờ", "(30/07 16:15)" in full2)
kiem("[12] tối: vẫn giữ tin quét thường", "Tin thế giới quét được" in full2)
kiem("[13] tối: đánh dấu đã gộp đúng id", danhdau_goi == [2])

# ----------- PHẢI CHẶN: dòng CHƯA tóm tắt KHÔNG được vào file (Huy chốt 01/08/2026)
# Trước bản vá, dòng này bị dán nguyên văn vào giữa bản tin, cắt ở 50.000 ký tự. Lô thật
# 01/08 dài 34.525 ký tự — bằng cả bốn mục tin cộng lại, không tiêu đề, không link.
danhdau_goi.clear()
DAI = "Câu văn dài lặp lại trong file gốc Jay Lâm gửi. " * 900


def doc_chua_xu_ly(now=None):
    return [dong_jaylam(6, DAI, "2026-07-30T09:00:00Z")], []


full6, err6 = chay(TOI, doc_chua_xu_ly, danhdau_ghi)
kiem("[14] chưa tóm tắt: KHÔNG có một chữ nguyên văn nào trong file",
     "Câu văn dài lặp lại" not in full6)
kiem("[52] chưa tóm tắt: KHÔNG dựng mục 'Tin Jay Lâm gửi'",
     "Tin Jay Lâm gửi" not in full6)
kiem("[15] chưa tóm tắt: CÓ cảnh báo để không hỏng câm",
     "CHƯA qua bước tóm tắt" in err6 and "id=6" in err6)
# Ca chính của bản vá 30/07/2026, giữ nguyên ý nghĩa sau khi bỏ nhánh nguyên văn: dòng chưa
# tóm tắt mà bị đóng sổ là mất hẳn cơ hội trở lại dưới dạng tin chuẩn.
kiem("[46] chưa tóm tắt: KHÔNG đánh dấu đã gộp (để bản tin sau gộp lại đủ)",
     danhdau_goi == [])

# --- ĐỐI CHỨNG chống lọc oan: cờ `da_xu_ly` bật mà `tom_tat` RỖNG vẫn là chưa xong.
# Đây là ca hỏng câm thật: `--ghi` đặt cờ theo lô nên một mục ghi thiếu tóm tắt sẽ lọt, và
# nếu chỉ xét cờ thì file in ra đúng một gạch đầu dòng trống.
danhdau_goi.clear()


def doc_co_bat_ma_rong(now=None):
    return [dong_jaylam(25, "Nguyên văn dòng cờ bật mà tóm tắt rỗng",
                        "2026-07-30T09:00:00Z", da_xu_ly=True, tieu_de="Có tiêu đề",
                        tom_tat="   ", nguon_ten="Reuters",
                        nguon_url="https://reuters.com/x-25")], []


full25, err25 = chay(TOI, doc_co_bat_ma_rong, danhdau_ghi)
kiem("[53] cờ bật mà tóm tắt rỗng: xử như CHƯA tóm tắt (không vào file, không đóng sổ)",
     "Tin Jay Lâm gửi" not in full25 and danhdau_goi == []
     and "CHƯA qua bước tóm tắt" in err25)

# --- đối chứng chống NỚI TAY: chỉ dòng CHƯA tóm tắt mới được tha, dòng đã xong vẫn đóng sổ
danhdau_goi.clear()


def doc_hon_hop(now=None):
    return [dong_jaylam(20, "Nguyên văn dòng chưa ai tóm tắt", "2026-07-30T09:00:00Z"),
            dong_xong(
                21, "Philippines nộp hồ sơ thềm lục địa mở rộng lên Liên Hợp Quốc",
                "2026-07-30T09:30:00Z",
                tom_tat="Philippines nộp lên Ủy ban Ranh giới Thềm lục địa hồ sơ một phần "
                        "về khu vực Tây Palawan, theo UNCLOS.",
                nguon_ten="Rappler", nguon_url="https://rappler.com/bai-cu-the-777")], []


full20, _ = chay(TOI, doc_hon_hop, danhdau_ghi)
kiem("[47] hỗn hợp: chỉ đóng sổ dòng ĐÃ tóm tắt, tha dòng chưa",
     danhdau_goi == [21])
kiem("[48] hỗn hợp: chỉ dòng ĐÃ tóm tắt vào file, dòng chưa bị giữ lại",
     "Tây Palawan" in full20 and "Nguyên văn dòng chưa ai tóm tắt" not in full20)

# --- chốt chặn LỚP HAI, đo bằng cách gọi THẲNG hàm (main đã lọc nên đi qua main thì lớp
# một che mất — mục 17 CLAUDE.md: bản hỏng gỡ một lớp mà lớp kia gánh thì ca vô dụng).
_d = Document()
_buf_tt = io.StringIO()
with contextlib.redirect_stderr(_buf_tt):
    MD.add_jaylam_item(_d, dong_jaylam(30, "Nguyên văn lọt qua lớp một",
                                       "2026-07-30T09:00:00Z"))
kiem("[54] add_jaylam_item nhận dòng chưa tóm tắt: BỎ QUA, không in đoạn nào + có kêu",
     len(_d.paragraphs) == 0 and "BỎ QUA" in _buf_tt.getvalue())

# --- ca MẤT TIN THẬT: hết khung ngày mà chưa phiên nào tóm tắt -> đóng sổ, phải kêu RIÊNG
danhdau_goi.clear()


def doc_qua_han_chua_tt(now=None):
    return [], [dong_jaylam(31, "Dòng chưa ai tóm tắt và đã quá hạn",
                            "2026-07-26T09:00:00Z"),
                dong_xong(32, "Tin đã tóm tắt nhưng quá hạn", "2026-07-26T09:00:00Z")]


_f, err31 = chay(TOI, doc_qua_han_chua_tt, danhdau_ghi)
# Dòng kêu riêng phải nêu ĐÚNG id chưa tóm tắt: gộp cả id=32 vào là mất khả năng phân biệt
# "tin quá hạn bình thường" với "tin chưa ai đọc mà đã đóng sổ".
_dong_kk = [l for l in err31.splitlines() if "CHƯA TỪNG được tóm tắt" in l]
kiem("[55] quá hạn mà CHƯA TỪNG tóm tắt: kêu riêng, nêu đúng id",
     len(_dong_kk) == 1 and "id=31" in _dong_kk[0] and "id=32" not in _dong_kk[0])
kiem("[56] quá hạn: cả hai đều đóng sổ (kể cả dòng chưa tóm tắt — hết đường chờ)",
     sorted(danhdau_goi) == [31, 32])

# ---- PHẢI CHẶN: tin Jay Lâm ĐÃ đi trong bản SÁNG cùng ngày (vá 01/08/2026)
# Lỗ đo thật tối 01/08: `loc_bo_tin_ca_sang` chỉ áp cho 03 mục quét thường, còn
# `loc_trung_jaylam` chỉ so tiêu đề với tin của CHÍNH bản đang dựng -> 04 tin Jay Lâm lặp
# nguyên si bản tin sáng cùng ngày, không lớp nào chặn.
URL_SANG = "https://reuters.com/da-di-ban-sang"
danhdau_goi.clear()


def chay_voi_so_sang(now, doc_fn, danhdau_fn, url_sang, data=None):
    """Như `chay()` nhưng ghim tập URL của ca sáng — sổ thật của repo không dựng được ở đây,
    và ca này đo phép LỌC chứ không đo đường đọc sổ (`test-so-da-gui.py` canh đường đó)."""
    goc = MD._url_ca_sang
    MD._url_ca_sang = lambda now_: set(url_sang)
    try:
        return chay(now, doc_fn, danhdau_fn, data=data)
    finally:
        MD._url_ca_sang = goc


def doc_trung_ban_sang(now=None):
    return [dong_xong(40, "Tin đã đi trong bản sáng nay", "2026-07-30T02:00:00Z",
                      tom_tat="Tóm tắt tin đã đi trong bản tin sáng nay rồi.",
                      nguon_url=URL_SANG),
            dong_xong(41, "Tin chỉ có trong bản tối", "2026-07-30T13:00:00Z",
                      tom_tat="Tóm tắt tin chỉ xuất hiện ở bản tin tối.")], []


full40, err40 = chay_voi_so_sang(TOI, doc_trung_ban_sang, danhdau_ghi, [URL_SANG])
kiem("[57] tin Jay trùng URL bản SÁNG: KHÔNG lặp lại ở bản tối",
     "đã đi trong bản tin sáng nay rồi" not in full40)
kiem("[58] tin Jay không trùng: vẫn vào bản tối (chống lọc oan)",
     "chỉ xuất hiện ở bản tin tối" in full40)
kiem("[59] tin Jay bị lọc vì trùng bản sáng: VẪN đóng sổ + có kêu",
     sorted(danhdau_goi) == [40, 41] and "bản tin SÁNG nay" in err40 and "id=40" in err40)

# ĐỐI CHỨNG cho đường đọc sổ THẬT, dựng sổ giả nên TẤT ĐỊNH — đọc sổ thật của repo thì ca
# này tự tắt sau 7 ngày (sổ chỉ giữ `GIU_NGAY = 7`), tức bản hỏng lọt mà bảng vẫn xanh.
import so_da_gui as SDG                                                    # noqa: E402

_so_cu = SDG.SO
_d_so = pathlib.Path(tempfile.mkdtemp(prefix="jl-so-"))
SDG.SO = _d_so / "da-gui-email.json"
SDG.SO.write_text(json.dumps({"lan_gui": [
    {"buoi": "sang", "luc": "2026-07-30T07:41:40+07:00", "urls": [URL_SANG]}]}),
    encoding="utf-8")
try:
    _sang, _toi = MD._url_ca_sang(SANG), MD._url_ca_sang(TOI)
finally:
    SDG.SO = _so_cu
    shutil.rmtree(_d_so, ignore_errors=True)
kiem("[60] _url_ca_sang ở bản SÁNG trả rỗng (lọc ở đó là giết chính lô vừa nạp)",
     _sang == set())
kiem("[61] _url_ca_sang ở bản TỐI đọc đúng dòng `sang` cùng ngày", _toi == {URL_SANG})

# Đọc sổ hỏng -> fail-OPEN có tiếng: giữ nguyên tin (hướng lệch là LẶP một bản tin, không
# phải MẤT tin) nhưng phải kêu. Ném lỗi ở đây là làm cả file .docx không ra đời.
_goc_doc_so = SDG.doc_so


def _no_tung():
    raise OSError("sổ hỏng")


SDG.doc_so = _no_tung
_buf_so = io.StringIO()
try:
    with contextlib.redirect_stderr(_buf_so):
        # Bắt lỗi TẠI ĐÂY: bản hỏng fail-CLOSED ném ra ngoài, không bắt thì cả bộ test chết
        # bằng traceback và `--tu-kiem` đọc thành "ca vẫn xanh" thay vì "ca đỏ".
        try:
            _kq_hong = MD._url_ca_sang(TOI)
        except Exception as _e:                      # noqa: BLE001
            _kq_hong = f"NÉM LỖI: {_e}"
finally:
    SDG.doc_so = _goc_doc_so
kiem("[62] đọc sổ ca sáng HỎNG: trả rỗng (giữ tin) + có kêu, không ném lỗi",
     _kq_hong == set() and "Không đọc được sổ ca sáng" in _buf_so.getvalue())

# --------------------------- PHẢI CHẶN: chống trùng với tin quét thường ĐÃ CÓ
danhdau_goi.clear()


def doc_trung(now=None):
    return [dong_xong(3, "Tin thế giới quét được", "2026-07-30T10:00:00Z")], []


full3, err3 = chay(TOI, doc_trung, danhdau_ghi)
kiem("[16] chống trùng: KHÔNG hiện mục khi trùng tin đã có",
     "Tin Jay Lâm gửi" not in full3)
kiem("[17] chống trùng: VẪN đánh dấu đã gộp (không nằm lại vĩnh viễn)",
     danhdau_goi == [3])
kiem("[18] chống trùng: có in cảnh báo lý do", "nghi trùng" in err3)

# --- chống trùng dựa trên `tieu_de` do phiên quét viết (không phải dòng đầu nguyên văn)
danhdau_goi.clear()


def doc_trung_theo_tieu_de(now=None):
    return [dong_jaylam(
        7, "Dòng đầu nguyên văn hoàn toàn khác hẳn không trùng gì cả",
        "2026-07-30T10:00:00Z", da_xu_ly=True,
        tieu_de="Tin thế giới quét được",
        tom_tat="Tóm tắt đủ dài cho guardrail của bước ghi, nội dung không quan trọng ở đây.",
        nguon_ten="Reuters", nguon_url="https://reuters.com/x-1")], []


full7, err7 = chay(TOI, doc_trung_theo_tieu_de, danhdau_ghi)
kiem("[19] chống trùng dùng `tieu_de` đã xử lý, không dùng dòng đầu nguyên văn",
     "Tin Jay Lâm gửi" not in full7 and danhdau_goi == [7])

# ------------- chống trùng GIỮA các dòng Jay Lâm với nhau, giữ dòng ĐẦU
danhdau_goi.clear()


def doc_hai_dong_trung_nhau(now=None):
    return [dong_xong(4, "Sự kiện đặc biệt hôm nay tại Hà Nội", "2026-07-30T08:00:00Z",
                      tom_tat="Sự kiện đặc biệt hôm nay tại Hà Nội, bản của dòng thứ nhất."),
            dong_xong(5, "Sự kiện đặc biệt hôm nay tại Hà Nội", "2026-07-30T08:05:00Z",
                      tom_tat="Sự kiện đặc biệt hôm nay tại Hà Nội, bản của dòng thứ hai.")
            ], []


full4, _ = chay(TOI, doc_hai_dong_trung_nhau, danhdau_ghi)
kiem("[20] trùng nội bộ: chỉ hiện MỘT lần trong file",
     full4.count("Sự kiện đặc biệt hôm nay tại Hà Nội") == 1)
# Cả hai đã tóm tắt nên cả hai đều đóng sổ: dòng 4 vì đã vào file, dòng 5 vì nội dung của nó
# coi như đã có mặt qua dòng 4 (xem docstring `loc_trung_jaylam`).
kiem("[21] trùng nội bộ: đóng sổ CẢ HAI, kể cả dòng bị lọc",
     sorted(danhdau_goi) == [4, 5])

# --------- KHUNG NGÀY THEO CHỦ ĐỀ (Huy chỉ ra 30/07: CNQS Mỹ nới 3 ngày lùi)
# `da_xu_ly=True` + không phải CNQS -> khung hẹp 2 ngày
XL = {"da_xu_ly": True}
kiem("[22] jaylam_qua_han: tin thường hôm nay -> False",
     MD.jaylam_qua_han(dict(XL, created_at="2026-07-30T02:00:00Z"), TOI) is False)
kiem("[23] jaylam_qua_han: tin thường HÔM QUA -> False (trong khung 2 ngày)",
     MD.jaylam_qua_han(dict(XL, created_at="2026-07-29T02:00:00Z"), TOI) is False)
kiem("[24] jaylam_qua_han: tin thường 2 ngày trước -> True",
     MD.jaylam_qua_han(dict(XL, created_at="2026-07-28T02:00:00Z"), TOI) is True)
kiem("[25] jaylam_qua_han: created_at HỎNG -> True (lỗi trả về phía KÊU)",
     MD.jaylam_qua_han(dict(XL, created_at="khong-phai-ngay"), TOI) is True)
kiem("[26] jaylam_qua_han: thiếu created_at -> True",
     MD.jaylam_qua_han(dict(XL), TOI) is True)

# PHẢI CHẶN chiều NGƯỢC: khung hẹp KHÔNG được loại oan tin CNQS Mỹ.
# Hôm nay 30/07 thì tin CNQS ngày 27/07 (3 ngày lùi) vẫn phải được giữ — đúng ví dụ Huy nêu.
CNQS = {"da_xu_ly": True, "la_cnqs": True}
kiem("[27] CNQS: tin 2 ngày trước -> GIỮ (khung thường đã loại)",
     MD.jaylam_qua_han(dict(CNQS, created_at="2026-07-28T02:00:00Z"), TOI) is False)
kiem("[28] CNQS: tin 3 ngày trước -> GIỮ (đúng ví dụ 27 giữ tin ngày 24)",
     MD.jaylam_qua_han(dict(CNQS, created_at="2026-07-27T02:00:00Z"), TOI) is False)
kiem("[29] CNQS: tin 4 ngày trước -> True (khung nới cũng có giới hạn)",
     MD.jaylam_qua_han(dict(CNQS, created_at="2026-07-26T02:00:00Z"), TOI) is True)
kiem("[30] CHƯA xử lý: hưởng khung RỘNG, không bị loại oan khi phiên quét chết",
     MD.jaylam_qua_han({"created_at": "2026-07-27T02:00:00Z"}, TOI) is False)
kiem("[31] jaylam_gioi_han_ngay: khớp hằng số của add_news/harvest (1 và 3)",
     (MD.JAYLAM_MAX_AGE_DAYS, MD.JAYLAM_MAX_AGE_DAYS_CNQS) == (1, 3)
     and MD.jaylam_gioi_han_ngay(dict(XL)) == 1
     and MD.jaylam_gioi_han_ngay(dict(CNQS)) == 3)

(trong, ngoai), err_loc = doc_that(
    [dong_jaylam(10, "Tin gửi hôm nay", "2026-07-30T03:00:00Z", da_xu_ly=True),
     dong_jaylam(11, "Tin gửi hôm qua", "2026-07-29T03:00:00Z", da_xu_ly=True),
     dong_jaylam(12, "Tin gửi ba ngày trước", "2026-07-27T03:00:00Z", da_xu_ly=True),
     dong_jaylam(13, "Tin CNQS ba ngày trước", "2026-07-27T03:00:00Z",
                 da_xu_ly=True, la_cnqs=True)], TOI)
kiem("[32] doc_tin_jaylam_chua_gop: tách đúng trong-khung / quá-hạn (CNQS được giữ)",
     [r["id"] for r in trong] == [10, 11, 13] and [r["id"] for r in ngoai] == [12])
kiem("[33] doc_tin_jaylam_chua_gop: cảnh báo nêu id bị bỏ kèm khung đã áp",
     "id=12" in err_loc and "khung 2 ngày" in err_loc)
kiem("[34] câu select đọc đủ cột do phiên quét ghi",
     all(c in open(GS / "make_docx.py", encoding="utf-8").read()
         for c in ("tieu_de", "tom_tat", "nguon_ten", "nguon_url", "da_xu_ly", "la_cnqs")))

# PHẢI CHẶN: tin quá hạn KHÔNG vào file nhưng VẪN được đánh dấu đã gộp
danhdau_goi.clear()


def doc_co_qua_han(now=None):
    return ([dong_xong(20, "Tin trong khung ngày", "2026-07-30T09:00:00Z",
                       tom_tat="Tóm tắt tin trong khung ngày, đủ dài để in ra.")],
            [dong_xong(21, "Tin quá hạn ba ngày", "2026-07-27T09:00:00Z",
                       tom_tat="Tóm tắt tin QUÁ HẠN ba ngày, đủ dài để in ra.")])


full8, _ = chay(TOI, doc_co_qua_han, danhdau_ghi)
kiem("[35] quá hạn: KHÔNG có trong file", "QUÁ HẠN ba ngày" not in full8)
kiem("[36] quá hạn: tin trong khung vẫn vào file",
     "Tóm tắt tin trong khung ngày" in full8)
kiem("[37] quá hạn: VẪN đánh dấu đã gộp (không nằm lại vĩnh viễn)",
     sorted(danhdau_goi) == [20, 21])

# PHẢI CHẶN: 0 tin quét + 0 tin hiển thị mà CÓ tin quá hạn -> vẫn đánh dấu
danhdau_goi.clear()
DATA_RONG = {"generatedAt": "2026-07-30", "worldNews": [], "usNews": [], "exercises": []}


def doc_chi_qua_han(now=None):
    return [], [dong_jaylam(22, "Chỉ có tin quá hạn", "2026-07-26T09:00:00Z")]


full9, _ = chay(TOI, doc_chi_qua_han, danhdau_ghi, data=DATA_RONG)
kiem("[38] 0 tin: không dựng file (DOCX rỗng)", full9 == "")
kiem("[39] 0 tin: tin quá hạn VẪN được đánh dấu đã gộp", danhdau_goi == [22])

# ------------------------------------------- KHÔNG trần số lượng (Huy chốt)
danhdau_goi.clear()


# ⚠️ Nội dung 12 dòng phải khác nhau THẬT SỰ, không phải khác một con số: bộ lọc chống trùng
# dùng Jaccard ≥ 0.6 trên token nên "Tin số 1"/"Tin số 2" bị coi là trùng — và đó là hành vi
# ĐÚNG của cổng. Dựng dữ liệu gần giống nhau là đo nhầm cổng (mục 17 CLAUDE.md).
NOI_DUNG_12 = [
    "Hạ viện Mỹ thông qua ngân sách quốc phòng năm tài khóa hai nghìn hai bảy",
    "Úc ký hợp đồng đóng tàu ngầm hạt nhân tại xưởng Osborne",
    "Trung Quốc điều tàu khảo sát vào vùng đặc quyền kinh tế Malaysia",
    "Nhật Bản nâng ngân sách phòng thủ tên lửa lên mức kỷ lục",
    "Philippines mở thêm căn cứ cho lực lượng luân phiên của đồng minh",
    "Lầu Năm Góc công bố hợp đồng bảo dưỡng động cơ máy bay vận tải",
    "Pháp điều tàu sân bay tới Ấn Độ Dương trong đợt triển khai mùa hè",
    "Ấn Độ thử nghiệm thành công tên lửa đạn đạo phóng từ tàu ngầm",
    "Hàn Quốc siết quy định xuất khẩu vật liệu bán dẫn sang nước thứ ba",
    "Indonesia đàm phán mua máy bay tuần thám biển tầm xa",
    "Anh triển khai khinh hạm tới khu vực Vịnh Ba Tư hộ tống thương thuyền",
    "Canada rót kinh phí hiện đại hóa hệ thống radar cảnh báo sớm Bắc Cực",
]


def doc_muoi_hai_tin(now=None):
    return [dong_xong(100 + i, NOI_DUNG_12[i], "2026-07-30T09:00:00Z",
                      tom_tat=NOI_DUNG_12[i] + ", theo nguồn tin của hãng.")
            for i in range(12)], []


full10, _ = chay(TOI, doc_muoi_hai_tin, danhdau_ghi)
thieu = [t[:30] for t in NOI_DUNG_12 if t[:40] not in full10]
kiem(f"[40] KHÔNG trần: cả 12 tin đều vào file (thiếu: {thieu})", not thieu)

# Vế đóng sổ đo phía HÀNG CHỜ, dựng riêng bằng 12 dòng CHƯA tóm tắt: phải tha nguyên lô,
# không sót một id nào. Ghép chung vào ca [40] thì một trong hai vế mất neo.
danhdau_goi.clear()


def doc_muoi_hai_chua_tt(now=None):
    return [dong_jaylam(200 + i, NOI_DUNG_12[i], "2026-07-30T09:00:00Z")
            for i in range(12)], []


chay(TOI, doc_muoi_hai_chua_tt, danhdau_ghi)
kiem("[41] 12 dòng chưa tóm tắt: tha CẢ LÔ, không đóng sổ id nào", danhdau_goi == [])

# ------------- PHẢI CHẶN: thiếu Supabase key -> ([], []) êm, KHÔNG vỡ file
os.environ.pop("SUPABASE_ANON_KEY", None)
os.environ.pop("DT_BOT_KEY", None)
full5, err5 = chay(TOI, MD.doc_tin_jaylam_chua_gop, lambda ids: None)
kiem("[42] thiếu SUPABASE_ANON_KEY/DT_BOT_KEY: không có mục Jay Lâm, không lỗi",
     "Tin Jay Lâm gửi" not in full5)
kiem("[43] thiếu key: có cảnh báo rõ ràng", "Thiếu SUPABASE_ANON_KEY" in err5)

# ------------------------------------------------------------------ TỰ KIỂM
BAN_HONG = [
    ("gỡ phép lọc khung ngày (trả nguyên rows)",
     "    trong, ngoai = [], []\n"
     "    for r in rows:\n"
     "        (ngoai if jaylam_qua_han(r, now) else trong).append(r)",
     "    trong, ngoai = list(rows), []"),
    ("bỏ khung NỚI của CNQS Mỹ -> mọi tin dùng khung hẹp (loại oan tin khí tài)",
     "    return JAYLAM_MAX_AGE_DAYS_CNQS if row.get(\"la_cnqs\") else JAYLAM_MAX_AGE_DAYS",
     "    return JAYLAM_MAX_AGE_DAYS"),
    ("dòng CHƯA xử lý mất khung rộng -> phiên quét chết là loại oan",
     '    if not row.get("da_xu_ly"):\n'
     "        return JAYLAM_MAX_AGE_DAYS_CNQS",
     '    if not row.get("da_xu_ly"):\n'
     "        return JAYLAM_MAX_AGE_DAYS"),
    ("nới khung hẹp thành 3 ngày -> cổng khung ngày mất răng với tin thường",
     "JAYLAM_MAX_AGE_DAYS = 1\nJAYLAM_MAX_AGE_DAYS_CNQS = 3",
     "JAYLAM_MAX_AGE_DAYS = 3\nJAYLAM_MAX_AGE_DAYS_CNQS = 3"),
    ("jaylam_qua_han fail-open khi created_at hỏng",
     "    except (ValueError, AttributeError, TypeError):\n"
     "        return True\n"
     "    return (now.date() - t.date()).days > jaylam_gioi_han_ngay(row)",
     "    except (ValueError, AttributeError, TypeError):\n"
     "        return False\n"
     "    return (now.date() - t.date()).days > jaylam_gioi_han_ngay(row)"),
    # Dựng lại ĐÚNG hành vi trước 01/08/2026 ở phía LỌC: dòng chưa tóm tắt lại lọt vào hàng
    # in và bị đóng sổ. Không in ra được nữa (chốt lớp hai chặn) nhưng đóng sổ là đủ mất tin.
    ("bỏ phép tách chưa-tóm-tắt -> dòng chưa xong lại lọt vào hàng in + bị đóng sổ",
     "    du, cho = [], []\n"
     "    for r in rows:\n"
     "        (du if da_tom_tat(r) else cho).append(r)",
     "    du, cho = list(rows), []"),
    ("da_tom_tat chỉ xét cờ, bỏ phép xét tóm tắt rỗng",
     '    return bool(row.get("da_xu_ly")) and bool((row.get("tom_tat") or "").strip())',
     '    return bool(row.get("da_xu_ly"))'),
    ("gỡ chốt chặn lớp hai trong add_jaylam_item",
     "    if not da_tom_tat(row):",
     "    if False:"),
    ("bỏ cảnh báo dòng chưa tóm tắt -> giữ lại trong im lặng",
     "    if jaylam_cho:\n"
     '        print(f"Tin Jay Lâm gửi: GIỮ LẠI {len(jaylam_cho)} dòng CHƯA qua bước tóm tắt "',
     "    if False:\n"
     '        print(f"Tin Jay Lâm gửi: GIỮ LẠI {len(jaylam_cho)} dòng CHƯA qua bước tóm tắt "'),
    ("bỏ tiếng kêu riêng cho dòng quá hạn CHƯA TỪNG tóm tắt (mất tin trong im lặng)",
     "    qh_chua_tt = [r for r in jaylam_qh if not da_tom_tat(r)]",
     "    qh_chua_tt = []"),
    # Dựng lại lỗ đo thật tối 01/08: tin Jay Lâm lặp nguyên si bản tin sáng cùng ngày.
    ("bỏ phép lọc tin Jay đã đi ở bản SÁNG -> lặp lại nguyên si bản sáng",
     "    jaylam_du, jaylam_sang = loc_jaylam_ca_sang(jaylam_du, now)",
     "    jaylam_sang = []"),
    ("bỏ nhóm trùng bản SÁNG khỏi đánh dấu -> nằm lại hàng chờ, tối nào cũng lọc lại",
     '    can_danh_dau = [r["id"] for r in jaylam_du + jaylam_sang + jaylam_qh]',
     '    can_danh_dau = [r["id"] for r in jaylam_du + jaylam_qh]'),
    ("_url_ca_sang áp cả bản SÁNG -> giết chính lô vừa nạp",
     "    if not la_buoi_toi(now):\n        return set()",
     "    if False:\n        return set()"),
    ("_url_ca_sang fail-CLOSED khi đọc sổ hỏng -> mất tin thay vì lặp tin",
     '        print(f"Không đọc được sổ ca sáng ({e}) — giữ nguyên toàn bộ tin.",'
     " file=sys.stderr)\n        return set()",
     '        print(f"Không đọc được sổ ca sáng ({e}) — giữ nguyên toàn bộ tin.",'
     ' file=sys.stderr)\n        raise'),
    ("bỏ dòng nhãn xác minh",
     "        set_font(pn.add_run(JAYLAM_NHAN_XAC_MINH), size=SIZE - 1, italic=True)",
     "        pass"),
    ("nhãn thời gian quay về CHỈ ghi giờ",
     '                .astimezone(VN).strftime("%d/%m %H:%M"))',
     '                .astimezone(VN).strftime("%H:%M"))'),
    ("chỉ đánh dấu dòng đã vào file, bỏ nhóm quá hạn -> hàng chờ nằm lại vĩnh viễn",
     '    can_danh_dau = [r["id"] for r in jaylam_du + jaylam_sang + jaylam_qh]',
     '    can_danh_dau = [r["id"] for r in jaylam_du + jaylam_sang]'),
    ("_jaylam_tieu_de bỏ ưu tiên `tieu_de` đã xử lý",
     '    td = (row.get("tieu_de") or "").strip()\n'
     "    if td:\n"
     "        return td[:200]",
     "    td = \"\""),
    ("nhánh 0 tin bỏ việc đánh dấu nhóm quá hạn",
     "        if jaylam_qh or jaylam_sang:\n"
     '            danh_dau_da_gop_jaylam([r["id"] for r in jaylam_sang + jaylam_qh])',
     "        pass"),
    # Dựng lại ĐÚNG hành vi trước 30/07/2026: khoá mục 5 vào riêng bản tối. Ngưỡng cũ không
    # còn dấu vết nào trong mã nên ca 04/05/44 mất neo nếu thiếu bản hỏng này.
    ("khoá lại mục 5 vào riêng bản TỐI (bug tin gửi muộn không bao giờ tới tay)",
     "    jaylam_goc, jaylam_qh = doc_tin_jaylam_chua_gop(now)",
     "    jaylam_goc, jaylam_qh = (doc_tin_jaylam_chua_gop(now) if la_buoi_toi(now)\n"
     "                             else ([], []))"),
    # Bản hỏng dựng lại ĐÚNG hành vi cũ đã gây mất mục 5 tối 30/07/2026: đóng sổ cả dòng
    # chưa tóm tắt, tức nó không bao giờ trở lại dưới dạng tin chuẩn.
    ("đóng sổ cả dòng CHƯA tóm tắt -> tin mất hẳn khỏi mọi bản tin",
     '    can_danh_dau = [r["id"] for r in jaylam_du + jaylam_sang + jaylam_qh]',
     '    can_danh_dau = [r["id"] for r in jaylam_goc + jaylam_sang + jaylam_qh]'),
]

KHAI_DO = {
    "gỡ phép lọc khung ngày (trả nguyên rows)": ["32", "33"],
    "bỏ khung NỚI của CNQS Mỹ -> mọi tin dùng khung hẹp (loại oan tin khí tài)":
        ["27", "28", "31", "32"],
    "dòng CHƯA xử lý mất khung rộng -> phiên quét chết là loại oan": ["30"],
    "nới khung hẹp thành 3 ngày -> cổng khung ngày mất răng với tin thường":
        ["24", "31", "32"],
    "jaylam_qua_han fail-open khi created_at hỏng": ["25", "26"],
    "bỏ phép tách chưa-tóm-tắt -> dòng chưa xong lại lọt vào hàng in + bị đóng sổ":
        ["41", "46", "52", "53"],
    "da_tom_tat chỉ xét cờ, bỏ phép xét tóm tắt rỗng": ["53"],
    "gỡ chốt chặn lớp hai trong add_jaylam_item": ["54"],
    "bỏ cảnh báo dòng chưa tóm tắt -> giữ lại trong im lặng": ["15", "53"],
    "bỏ tiếng kêu riêng cho dòng quá hạn CHƯA TỪNG tóm tắt (mất tin trong im lặng)": ["55"],
    "bỏ dòng nhãn xác minh": ["10"],
    "nhãn thời gian quay về CHỈ ghi giờ": ["11"],
    "chỉ đánh dấu dòng đã vào file, bỏ nhóm quá hạn -> hàng chờ nằm lại vĩnh viễn":
        ["37", "56"],
    "_jaylam_tieu_de bỏ ưu tiên `tieu_de` đã xử lý": ["19"],
    "nhánh 0 tin bỏ việc đánh dấu nhóm quá hạn": ["39"],
    # 45 đỏ theo là ĐÚNG: khoá lại bản sáng thì tin quá hạn cũng không được đọc ra để đóng
    # sổ, tức nó nằm lại hàng chờ — chính vế thứ hai của cùng một bug.
    "khoá lại mục 5 vào riêng bản TỐI (bug tin gửi muộn không bao giờ tới tay)":
        ["04", "05", "44", "45"],
    "đóng sổ cả dòng CHƯA tóm tắt -> tin mất hẳn khỏi mọi bản tin": ["41", "46", "47"],
    "bỏ phép lọc tin Jay đã đi ở bản SÁNG -> lặp lại nguyên si bản sáng": ["57"],
    "bỏ nhóm trùng bản SÁNG khỏi đánh dấu -> nằm lại hàng chờ, tối nào cũng lọc lại": ["59"],
    "_url_ca_sang áp cả bản SÁNG -> giết chính lô vừa nạp": ["60"],
    "_url_ca_sang fail-CLOSED khi đọc sổ hỏng -> mất tin thay vì lặp tin": ["62"],
}


def _so_ca(ten):
    return ten[1:3] if ten.startswith("[") else ""


def tu_kiem():
    """Dựng từng bản `make_docx.py` đã GỠ một lớp bảo vệ rồi chạy lại chính bộ ca này.

    Bản hỏng nằm trong thư mục copy riêng mang PID + sha1 NỘI DUNG (mục 17 + 23 CLAUDE.md):
    tên cố định làm hai phiên chạy chồng xoá bản hỏng của nhau, và cùng đường dẫn làm
    `__pycache__` phát lại bytecode của bản hỏng TRƯỚC.
    """
    goc = (GS / "make_docx.py").read_text(encoding="utf-8")
    tong, trot = 0, []
    for ten, tim, thay in BAN_HONG:
        tong += 1
        if goc.count(tim) != 1:
            trot.append(f"{ten}: chuỗi neo khớp {goc.count(tim)} chỗ (phải đúng 1)")
            continue
        hong = goc.replace(tim, thay)
        sha = hashlib.sha1(hong.encode("utf-8")).hexdigest()[:8]
        d = pathlib.Path(tempfile.mkdtemp(prefix=f"jl-hong-{os.getpid()}-{sha}-"))
        try:
            for f in GS.glob("*.py"):
                shutil.copy2(f, d / f.name)
            (d / "make_docx.py").write_text(hong, encoding="utf-8")
            env = dict(os.environ, MAKEDOCX_DIR=str(d))
            p = SP.run([sys.executable, str(pathlib.Path(__file__).resolve())],
                       capture_output=True, text=True, env=env, timeout=300)
            do = {_so_ca(l[2:]) for l in p.stdout.splitlines() if l.startswith("✗ ")}
            can = set(KHAI_DO.get(ten, []))
            if p.returncode == 0:
                trot.append(f"{ten}: bộ test VẪN XANH -> không bắt được lỗi")
            elif len(do) >= len([l for l in p.stdout.splitlines()
                                 if l[:1] in ("✓", "✗")]):
                trot.append(f"{ten}: ĐỎ TOÀN BỘ ca -> phép thay phá cú pháp, không gỡ lớp vá")
            elif not can & do:
                trot.append(f"{ten}: ca cần đỏ {sorted(can)} vẫn xanh; đỏ thực tế "
                            f"{sorted(do)}")
            else:
                print(f"  ✓ {ten}: bắt được (ca đỏ {sorted(do)})")
        finally:
            shutil.rmtree(d, ignore_errors=True)
    print()
    if trot:
        print(f"TRƯỢT {len(trot)}/{tong} bản hỏng:")
        for t in trot:
            print("  - " + t)
        return 1
    print(f"✅ {tong}/{tong} bản hỏng đều bị bộ test bắt.")
    return 0


so_dat = sum(1 for _, ok in CA if ok)
print(f"\n{so_dat}/{len(CA)} ca đạt")
if "--tu-kiem" in sys.argv:
    if so_dat != len(CA):
        print("Bản THẬT đã đỏ — sửa xong hãy chạy --tu-kiem.")
        sys.exit(1)
    print("\n=== TỰ KIỂM: dựng bản make_docx.py đã gỡ lớp bảo vệ ===")
    sys.exit(tu_kiem())
sys.exit(0 if so_dat == len(CA) else 1)
