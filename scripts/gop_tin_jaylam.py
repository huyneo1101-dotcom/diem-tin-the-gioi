#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gộp tin Jay Lâm gửi trong ngày (qua bot, file .docx) thành MỘT file .docx, gửi cho Huy.

    python3 scripts/gop_tin_jaylam.py

Đọc bảng `dt_jaylam_inbox` (ngay_vn = hôm nay giờ VN, da_gop = false) qua mã riêng `x-dt-key`
— giống cách `telegram_bot.py:lich_su_gan_day()` đọc `dt_bot_hoi`, KHÔNG phải service key.
Không có tin nào trong ngày -> thoát êm (mã 10), không gửi gì.

⚠️ Đây là TÀI LIỆU RIÊNG cho Huy đọc (Huy chốt 30/07/2026), KHÔNG phải bản tin công khai:
gửi qua `send_document` CHỈ tới `telegram_bot.chat_chu()` (chat riêng đầu tiên trong
TELEGRAM_CHAT_ID) — TUYỆT ĐỐI không lặp qua cả `TELEGRAM_CHAT_ID`, vì danh sách đó có CẢ Jay
Lâm (xem CLAUDE.md mục "Ràng buộc kênh — Jay Lâm là NGƯỜI NGOÀI"). Gửi cho cả danh sách là gửi
thẳng nội dung tổng hợp về chính người đã gửi tin.
"""
import datetime
import json
import os
import pathlib
import subprocess
import sys
import zoneinfo

sys.path.insert(0, str(pathlib.Path(__file__).resolve().parent))
from tg_api import call, send_document, kiem_cau_hinh  # noqa: E402
import telegram_bot  # noqa: E402 — tái dùng chats_cho_phep()/chat_chu()/_anon_key()/_dt_bot_key()

from docx import Document  # noqa: E402
from docx.shared import Pt, Inches  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

VN = zoneinfo.ZoneInfo("Asia/Ho_Chi_Minh")
SUPABASE_URL = "https://ltmlueqkajqmduoqghdf.supabase.co"
BANG = "dt_jaylam_inbox"
FONT = "Times New Roman"
SIZE = 13


def doc_ngay(key, dt_key, ngay):
    """SELECT các dòng chưa gộp của một ngày. Trả None nếu đọc hỏng (khác [] = rỗng thật)."""
    p = subprocess.run(
        ["curl", "-sS", "--max-time", "30",
         f"{SUPABASE_URL}/rest/v1/{BANG}"
         f"?select=id,chat_id,ten,ten_file,noi_dung,created_at"
         f"&ngay_vn=eq.{ngay}&da_gop=eq.false&order=created_at.asc",
         "-H", f"apikey: {key}", "-H", f"Authorization: Bearer {key}",
         "-H", f"x-dt-key: {dt_key}"],
        capture_output=True, text=True, timeout=35)
    try:
        rows = json.loads(p.stdout)
    except (json.JSONDecodeError, TypeError):
        print(f"Đọc Supabase hỏng: {p.stdout[:200]} {p.stderr[:200]}", file=sys.stderr)
        return None
    return rows if isinstance(rows, list) else None


def danh_dau_da_gop(key, dt_key, ids):
    """UPDATE da_gop=true — tránh gộp trùng nếu script chạy lại trong ngày."""
    if not ids:
        return
    subprocess.run(
        ["curl", "-sS", "--max-time", "30", "-X", "PATCH",
         f"{SUPABASE_URL}/rest/v1/{BANG}?id=in.({','.join(str(i) for i in ids)})",
         "-H", f"apikey: {key}", "-H", f"Authorization: Bearer {key}",
         "-H", f"x-dt-key: {dt_key}", "-H", "Content-Type: application/json",
         "-H", "Prefer: return=minimal", "-d", '{"da_gop": true}'],
        capture_output=True, text=True, timeout=35)


def gio_vn(iso):
    try:
        return (datetime.datetime.fromisoformat(iso.replace("Z", "+00:00"))
                .astimezone(VN).strftime("%H:%M"))
    except (ValueError, AttributeError):
        return "--:--"


def dung_docx(rows, ngay):
    d = Document()
    for section in d.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    style = d.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(SIZE)

    tieu = d.add_paragraph()
    tieu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tieu.add_run(f"TIN JAY LÂM GỬI NGÀY {ngay}")
    r.bold = True
    r.font.size = Pt(SIZE + 1)
    r.font.name = FONT

    for row in rows:
        meta = d.add_paragraph()
        mr = meta.add_run(f"{gio_vn(row.get('created_at') or '')} — "
                           f"{row.get('ten_file') or '(không tên)'} "
                           f"({row.get('ten') or 'Jay Lâm'})")
        mr.bold = True
        mr.font.name = FONT
        noi = d.add_paragraph(row.get("noi_dung") or "(rỗng)")
        for run in noi.runs:
            run.font.name = FONT
        d.add_paragraph("")

    path = f"/tmp/Tin-Jay-Lam-{ngay}.docx"
    d.save(path)
    return path


def main():
    token = os.environ.get("TELEGRAM_BOT_TOKEN", "").strip()
    chats = telegram_bot.danh_sach_chat()
    gate = kiem_cau_hinh(token, chats, "gop-tin-jaylam")
    if gate is not None:
        return gate

    key = telegram_bot._anon_key()
    dt_key = telegram_bot._dt_bot_key()
    if not key or not dt_key:
        print("Thiếu SUPABASE_ANON_KEY hoặc DT_BOT_KEY — không đọc được bảng dt_jaylam_inbox.",
              file=sys.stderr)
        return 1

    ngay = datetime.datetime.now(VN).date().isoformat()
    rows = doc_ngay(key, dt_key, ngay)
    if rows is None:
        return 1
    if not rows:
        print(f"Không có tin Jay Lâm gửi ngày {ngay}.")
        return 10

    chu = telegram_bot.chat_chu()
    if not chu:
        print("Không xác định được chat của Huy (TELEGRAM_CHAT_ID rỗng).", file=sys.stderr)
        return 1

    path = dung_docx(rows, ngay)
    r = send_document(token, chu, path,
                       caption=f"📎 {len(rows)} tin Jay Lâm gửi ngày {ngay}")
    if not r.get("ok"):
        print(f"Gửi file hỏng: {r.get('description')}", file=sys.stderr)
        return 1
    danh_dau_da_gop(key, dt_key, [row["id"] for row in rows])
    print(f"Đã gửi {len(rows)} tin tổng hợp ngày {ngay} cho Huy.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
